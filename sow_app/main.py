# app.py - COMPLETE SOW MANAGEMENT SYSTEM WITH SHAREPOINT INTEGRATION
from docxtpl import DocxTemplate
import streamlit as st
from datetime import datetime, date, timedelta
from io import BytesIO
import pandas as pd
import os
import warnings
import base64
import json
import requests
import time
import numpy as np
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# Suppress warnings
warnings.filterwarnings('ignore')

# Page configuration
st.set_page_config(
    page_title="SOW Generator",
    layout="wide",
    page_icon="📋",
    initial_sidebar_state="collapsed"  # Hide sidebar
)

# ============================================================================
# CUSTOM CSS STYLING
# ============================================================================
st.markdown("""
<style>
/* Hide default Streamlit elements */
header {visibility: hidden;}
footer {visibility: hidden;}
.stDeployButton {display: none;}
.block-container {padding-top: 0rem;}

/* Hide sidebar */
section[data-testid="stSidebar"] {display: none !important;}

/* Custom status badges */
.status-badge {
    padding: 4px 12px;
    border-radius: 20px;
    font-size: 12px;
    font-weight: 600;
    display: inline-block;
}
.status-draft { background: #e0e0e0; color: #616161; }
.status-pending { background: #fff3cd; color: #856404; }
.status-approved { background: #d4edda; color: #155724; }
.status-rejected { background: #f8d7da; color: #721c24; }

/* Custom button styles */
.stButton > button {
    border-radius: 8px;
    transition: all 0.3s ease;
}
.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 12px rgba(0,0,0,0.15);
}

/* Custom header styling */
.custom-header {
    background: linear-gradient(90deg, #0a0f1e, #13203d, #1f3d6d);
    padding: 20px;
    border-radius: 10px;
    margin-bottom: 20px;
    color: white;
}
.custom-header h1 {
    color: white;
    margin: 0;
}
.custom-header p {
    color: #b0c4de;
    margin: 5px 0 0 0;
}

/* Login form styling */
.login-container {
    max-width: 400px;
    margin: 100px auto;
    padding: 40px;
    background: white;
    border-radius: 15px;
    box-shadow: 0 10px 30px rgba(0,0,0,0.1);
}
.login-header {
    text-align: center;
    margin-bottom: 30px;
}
.login-header h2 {
    color: #1f3d6d;
    margin-bottom: 10px;
}
.login-header p {
    color: #666;
}

/* Table styling */
.dataframe {
    border-radius: 8px;
    overflow: hidden;
}

/* User info in header */
.user-info {
    position: absolute;
    top: 20px;
    right: 20px;
    background: rgba(255,255,255,0.1);
    padding: 8px 15px;
    border-radius: 20px;
    font-size: 14px;
}
.logout-btn {
    position: absolute;
    top: 20px;
    right: 20px;
    background: rgba(255,255,255,0.1);
    color: white;
    border: none;
    padding: 8px 15px;
    border-radius: 20px;
    font-size: 14px;
    cursor: pointer;
}
.logout-btn:hover {
    background: rgba(255,255,255,0.2);
}

/* View/Edit button styling */
.view-btn {
    background-color: #4CAF50 !important;
    color: white !important;
    border: none !important;
    padding: 5px 10px !important;
    border-radius: 4px !important;
    font-size: 12px !important;
    cursor: pointer !important;
}
.view-btn:hover {
    background-color: #45a049 !important;
}

/* Approval section styling */
.approval-section {
    background-color: #f8f9fa;
    padding: 20px;
    border-radius: 10px;
    margin-top: 20px;
    border-left: 5px solid #0a0f1e;
}
.approve-btn {
    background-color: #28a745 !important;
    color: white !important;
    border: none !important;
}
.reject-btn {
    background-color: #dc3545 !important;
    color: white !important;
    border: none !important;
}

/* Action buttons in table */
.action-cell {
    display: flex;
    gap: 5px;
}
</style>
""", unsafe_allow_html=True)

# ============================================================================
# CONFIGURATION
# ============================================================================
class Config:
    # ======== UPDATE THESE URLs WITH YOUR POWER AUTOMATE FLOW URLs ========
    POWER_AUTOMATE_URLS = {
        "save_record": "https://defaulted788079405e492bbc316ba6912792.09.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/53097801090042b2b7fa1f3f9bfac9e2/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=ZX_J0dUOdGpGt2DcmEAm_qFzLdPXfFv6fkzOsQvm-_w",
        "get_records": "https://defaulted788079405e492bbc316ba6912792.09.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/78733188f09f4ab19455d4f5ec755296/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=f4rvtDGpIxjXF7O1UGLK4KUtnsP7Sy_EuXswX6ZdU-o",
        "update_status": "https://defaulted788079405e492bbc316ba6912792.09.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/9f22dfef2dd0403fb780810cc3c54b45/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=EDsK9ooYU6bXx7QZNQhwlKPHXAMPMHpPihDQea225Ys",
        "upload_document": "https://defaulted788079405e492bbc316ba6912792.09.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/c5172488c3244b6aa59242f0562dbb10/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=MV-4UfPewDgKk1QuVS5-z0_yUVwAKQz4d_aP9S3qB0c",
        "get_document": "https://defaulted788079405e492bbc316ba6912792.09.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/4e3f51632a5d4db79976427be5a54b6b/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=3Rgaj1vUDDydVP8x7jEbrCzBDDZ_r_N7aKaYYnisAok",
        "get_sow_details": "https://defaulted788079405e492bbc316ba6912792.09.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/78733188f09f4ab19455d4f5ec755296/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=f4rvtDGpIxjXF7O1UGLK4KUtnsP7Sy_EuXswX6ZdU-o"
    }
    
    # SharePoint Document Libraries
    DOCUMENT_LIBRARIES = {
        "templates": "SOW_Templates",
        "generated": "Onboarding Details",
        "approved": "SOW_Approved",
        "folder_path": "SOWs"
    }
    
    # Template mapping
    TEMPLATE_MAPPING = {
        "Fixed Fee": "Fixed_Fee_Template.docx",
        "T&M": "T&M_Template.docx",
        "Change Order": "Change_Order_Template.docx"
    }
    
    # Login credentials
    LEGAL_USER_EMAIL = "legal@cloudlabsit.com"
    # COMMON_PASSWORD = "CloudLabs@123"  # Common password for all users
    
    # Legal team emails (for approval dashboard access)
    LEGAL_TEAM = ["legal@cloudlabsit.com", "admin@cloudlabsit.com"]
    
    # Status values
    STATUS_DRAFT = "Draft"
    STATUS_PENDING = "Pending Review"
    STATUS_APPROVED = "Approved"
    STATUS_REJECTED = "Rejected"
    
    # SharePoint List Name
    SHAREPOINT_LIST = "SOW_Records"

# ============================================================================
# SHAREPOINT SERVICE VIA POWER AUTOMATE
# ============================================================================
class SharePointService:
    # Add this to the SharePointService class

    def update_sow_record(self, item_id, sow_data):
        """Update existing SOW record in SharePoint"""
        try:
            payload = {
                "operation": "update_item",
                "list_name": self.config.SHAREPOINT_LIST,
                "item_id": item_id,
                "updates": sow_data
            }
            
            print(f"🔍 DEBUG: Updating SOW record {item_id}")
            print(f"Update data: {json.dumps(sow_data, default=str, indent=2)}")
            
            result = self._call_power_automate("update_status", payload)  # Reusing update_status flow
            
            if result:
                return {
                    "success": True,
                    "data": result,
                    "message": "SOW record updated successfully"
                }
            return {
                "success": False,
                    "message": "Failed to update SOW record"
                }
        except Exception as e:
            print(f"❌ Error updating SOW record: {str(e)}")
            return {
                "success": False,
                "message": f"Error: {str(e)}"
            }
    
    def __init__(self):
        self.config = Config()
    
    def _call_power_automate(self, flow_name, payload=None):
        """Call Power Automate flow - FIXED VERSION"""
        try:
            url = self.config.POWER_AUTOMATE_URLS.get(flow_name)
            if not url:
                print(f"❌ No URL for flow: {flow_name}")
                return None
            
            headers = {"Content-Type": "application/json"}
            
            print(f"🔍 DEBUG: Calling {flow_name} flow")
            print(f"🔍 DEBUG: URL: {url}")
            
            if payload:
                # CRITICAL FIX: Use ensure_ascii=False to preserve binary data
                json_payload = json.dumps(payload, ensure_ascii=False)
                
                # Debug: Check if file_content is in the payload
                if 'file_content' in payload:
                    fc = payload['file_content']
                    print(f"🔍 DEBUG: file_content in payload:")
                    print(f"  - Type: {type(fc)}")
                    print(f"  - Length: {len(fc) if fc else 0}")
                    print(f"  - First 50 chars: {fc[:50] if fc else 'EMPTY'}")
                    print(f"  - Is string: {isinstance(fc, str)}")
                    print(f"  - Not empty: {bool(fc)}")
                else:
                    print("❌ ERROR: file_content NOT in payload!")
                
                print(f"🔍 DEBUG: Sending JSON payload ({len(json_payload)} chars)")
                
                try:
                    response = requests.post(url, data=json_payload.encode('utf-8'), headers=headers, timeout=30)
                except Exception as req_error:
                    print(f"❌ Request error: {req_error}")
                    return None
            else:
                response = requests.post(url, headers=headers, timeout=30)
            
            print(f"🔍 DEBUG: Response status: {response.status_code}")
            
            if response.status_code != 200:
                print(f"❌ ERROR Response: {response.text[:500]}")
            
            response.raise_for_status()
            
            try:
                return response.json()
            except:
                return {"success": True, "raw": response.text}
                
        except Exception as e:
            print(f"❌ Error in _call_power_automate: {str(e)}")
            return None
    
    def save_sow_record(self, sow_data):
        """Save SOW record to SharePoint list - MATCHES POWER AUTOMATE SCHEMA"""
        
        # Create payload that matches Power Automate schema
        payload = {
            "operation": "create_sow_record",  # Must be exactly this
            "sow_data": sow_data
        }
        
        # Debug: Show what we're sending
        print("=== Sending to Power Automate ===")
        print(json.dumps(payload, default=str, indent=2))
        
        result = self._call_power_automate("save_record", payload)
        
        if result:
            return {
                "success": True,
                "data": result,
                "message": "SOW record saved successfully"
            }
        return {
            "success": False,
            "message": "Failed to save SOW record"
        }
    
    def get_sow_records(self, status=None, status_filter=None, user_filter=None, 
               client_filter=None, project_type_filter=None,
               date_from=None, date_to=None):
        """Get SOW records from SharePoint"""
        
        # Handle both 'status' and 'status_filter' for backwards compatibility
        if status is not None:
            status_filter = status
        
        payload = {
            "operation": "get_items",
            "list_name": self.config.SHAREPOINT_LIST,
            "filters": {}
        }
        
        # Build filters
        filters = {}
        if status_filter:
            filters["status"] = status_filter
        if user_filter:
            filters["created_by"] = user_filter
        if client_filter:
            filters["client"] = client_filter
        if project_type_filter:
            filters["project_type"] = project_type_filter
        if date_from:
            filters["date_from"] = date_from
        if date_to:
            filters["date_to"] = date_to
        
        payload["filters"] = filters
        
        # DEBUG: Print what we're sending
        print("=== GET RECORDS DEBUG ===")
        print(f"Payload: {json.dumps(payload, indent=2)}")
        
        result = self._call_power_automate("get_records", payload)
        
        # DEBUG: Print what we got back
        print(f"Result type: {type(result)}")
        print(f"Result: {result}")
        
        if result:
            print(f"Result keys: {result.keys() if isinstance(result, dict) else 'Not a dict'}")
            if isinstance(result, dict) and "items" in result:
                print(f"Items count: {len(result['items'])}")
                print(f"First item (if exists): {result['items'][0] if len(result['items']) > 0 else 'No items'}")
            else:
                print(f"❌ 'items' key not found in result. Available keys: {list(result.keys()) if isinstance(result, dict) else 'N/A'}")
        else:
            print("❌ Result is None or empty")
        
        if result and "items" in result:
            try:
                print(f"✅ Attempting to create DataFrame from {len(result['items'])} items...")
                df = pd.DataFrame(result["items"])
                print(f"✅ DataFrame created successfully with {len(df)} rows")
                print(f"DataFrame columns: {list(df.columns)}")
                return {
                    "success": True,
                    "data": df,
                    "count": len(df),
                    "message": "Records retrieved successfully"
                }
            except Exception as e:
                print(f"❌ DataFrame creation error: {str(e)}")
                import traceback
                traceback.print_exc()
                return {
                    "success": False,
                    "error": str(e),
                    "message": f"Failed to parse records: {str(e)}"
                }
        
        print("❌ No items in result or result is None")
        return {
            "success": False,
            "data": pd.DataFrame(),
            "message": "No records found or flow not configured"
        }
    
    def get_sow_by_id(self, item_id):
        """Get specific SOW by ID from SharePoint"""
        payload = {
            "operation": "get_item_by_id",
            "list_name": self.config.SHAREPOINT_LIST,
            "item_id": str(item_id)
        }
        
        result = self._call_power_automate("get_sow_details", payload)
        
        if result and "item" in result:
            return {
                "success": True,
                "data": result["item"],
                "message": "SOW retrieved successfully"
            }
        
        return {
            "success": False,
            "message": "Failed to retrieve SOW"
        }
    
    def update_sow_status(self, item_id, status, comments="", approver_email=""):
        """Update SOW status in SharePoint"""
        payload = {
            "operation": "update_item",
            "list_name": self.config.SHAREPOINT_LIST,
            "item_id": item_id,
            "updates": {
                "status": status,
                "approver_comments": comments,
                "approved_by": approver_email,
                "approval_date": datetime.now().isoformat() if status == Config.STATUS_APPROVED else ""
            }
        }
        
        result = self._call_power_automate("update_status", payload)
        if result:
            return {
                "success": True,
                "data": result,
                "message": "Status updated successfully"
            }
        return {
            "success": False,
            "message": "Failed to update status"
        }
    
    def upload_document(self, file_bytes, file_name, metadata):
        """Upload document to SharePoint - UPDATED FOR CORRECT PATH"""
        try:
            print(f"🔍 DEBUG upload_document called:")
            print(f"  - file_name: {file_name}")
            print(f"  - file_bytes length: {len(file_bytes) if file_bytes else 'EMPTY'}")
            
            if not file_bytes or len(file_bytes) == 0:
                print("❌ ERROR: file_bytes is empty!")
                return {"success": False, "message": "File bytes are empty"}
            
            # Convert to Base64
            try:
                if hasattr(file_bytes, 'read'):
                    file_bytes = file_bytes.read()
                
                file_base64 = base64.b64encode(file_bytes).decode('utf-8')
                print(f"✅ DEBUG: Base64 conversion successful ({len(file_base64)} chars)")
                
                if not file_base64 or len(file_base64) < 10:
                    raise ValueError("Base64 string too short or empty")
                    
            except Exception as conv_error:
                print(f"❌ Base64 conversion error: {conv_error}")
                return {
                    "success": False,
                    "error": str(conv_error),
                    "message": "Failed to convert file to Base64"
                }
            
            # Build payload - UPDATED library_name
            payload = {
                "operation": "upload_document",
                "library_name": "Onboarding Details",  # ✅ Changed
                "folder_path": "SOWs",  # ✅ Added folder
                "file_name": file_name,
                "file_content": file_base64,
                "metadata": {
                    "sow_number": metadata.get("sow_number", ""),
                    "created_by": metadata.get("created_by", ""),
                    "status": metadata.get("status", ""),
                    "project_type": metadata.get("project_type", ""),
                    "client": metadata.get("client", "")
                }
            }
            
            print(f"✅ DEBUG: Final payload ready")
            
            # Call Power Automate
            result = self._call_power_automate("upload_document", payload)
            
            if result:
                print(f"✅ SUCCESS: Power Automate returned: {result}")
                return {
                    "success": True,
                    "data": result,
                    "message": "Document uploaded successfully"
                }
            
            return {
                "success": False,
                "message": "No response from Power Automate"
            }
                
        except Exception as e:
            print(f"❌ FATAL ERROR in upload_document: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "message": f"Fatal error: {str(e)}"
            }
    
    def get_document(self, item_id=None, file_name=None, library_name=None):
        """Get document from SharePoint"""
        try:
            # Build payload based on available parameters
            if item_id:
                # Get document using item ID from list
                payload = {
                    "operation": "get_document",
                    "item_id": str(item_id)
                }
            elif file_name and library_name:
                # Get document directly by filename
                payload = {
                    "operation": "get_document",
                    "library_name": library_name,
                    "file_name": file_name
                }
            else:
                return None
            
            print(f"🔍 DEBUG: Getting document with payload: {payload}")
            
            result = self._call_power_automate("get_document", payload)
            
            if result and result.get("success"):
                # Decode Base64 content
                file_content_base64 = result.get("file_content")
                if file_content_base64:
                    file_bytes = base64.b64decode(file_content_base64)
                    return file_bytes
            
            return None
            
        except Exception as e:
            print(f"❌ Error getting document: {str(e)}")
            return None

# ============================================================================
# EXCEL EXPORTER CLASS
# ============================================================================
class ExcelExporter:
    """Class to handle Excel file creation for SOW data"""
    
    def __init__(self, output_folder="generated_excels"):
        self.output_folder = output_folder
        self.ensure_folder_exists()
    
    def ensure_folder_exists(self):
        """Create output folder if it doesn't exist"""
        os.makedirs(self.output_folder, exist_ok=True)
    
    def create_fixed_fee_milestone_excel(self, sow_data, milestone_df):
        """Create Excel file for Fixed Fee milestone payments - UPDATED WITH BETTER ERROR HANDLING"""
        try:
            # Extract data with better defaults
            sow_number = sow_data.get("sow_num", "UNKNOWN")
            sow_name = sow_data.get("sow_name", "Unknown SOW")
            client = sow_data.get("Client_Name", "Unknown Client")
            total_fees = float(sow_data.get("Fees_al", 0))
            
            # Handle date conversion
            start_date = sow_data.get("start_date", date.today())
            if isinstance(start_date, str):
                try:
                    start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
                except:
                    start_date = date.today()
            
            end_date = sow_data.get("end_date", date.today())
            if isinstance(end_date, str):
                try:
                    end_date = datetime.strptime(end_date, "%Y-%m-%d").date()
                except:
                    end_date = date.today()
            
            # Create workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Milestone Payments"
            
            # Header styling
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            thin_border = Border(left=Side(style='thin'), 
                                right=Side(style='thin'), 
                                top=Side(style='thin'), 
                                bottom=Side(style='thin'))
            
            # Write SOW Information
            ws['A1'] = f"SOW: {sow_number} - {sow_name}"
            ws['A1'].font = Font(bold=True, size=14, color="366092")
            ws.merge_cells('A1:E1')
            
            ws['A2'] = f"Client: {client}"
            ws['A2'].font = Font(bold=True)
            
            ws['A3'] = f"Total Contract Value: ${total_fees:,.2f}"
            ws['A3'].font = Font(bold=True)
            
            ws['A4'] = f"Contract Period: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}"
            
            # Blank row
            ws['A6'] = "Milestone Payment Schedule"
            ws['A6'].font = Font(bold=True, size=12)
            ws.merge_cells('A6:E6')
            
            # Column headers for milestones
            headers = ["Milestone #", "Services / Deliverables", "Due Date", 
                    "Payment Allocation (%)", "Payment Amount ($)"]
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=8, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border
                ws.column_dimensions[get_column_letter(col)].width = 25
            
            # Write milestone data
            if milestone_df is not None and not milestone_df.empty:
                # Ensure we have the right column names
                column_mapping = {
                    'milestone_no': ['milestone_no', 'Milestone #', 'Milestone No'],
                    'services': ['services', 'Services', 'Services / Deliverables'],
                    'due_date': ['due_date', 'Due Date', 'Milestone Due Date'],
                    'allocation': ['allocation', 'Allocation', 'Payment Allocation (%)'],
                    'net_pay': ['net_pay', 'Net Pay', 'Payment Amount ($)', 'Payment']
                }
                
                # Map columns if needed
                for target_col, possible_names in column_mapping.items():
                    if target_col not in milestone_df.columns:
                        for possible_name in possible_names:
                            if possible_name in milestone_df.columns:
                                milestone_df[target_col] = milestone_df[possible_name]
                                break
                
                # Convert date columns
                if 'due_date' in milestone_df.columns:
                    milestone_df['due_date'] = pd.to_datetime(milestone_df['due_date']).dt.date
                
                # Write data rows
                for idx, row in enumerate(milestone_df.itertuples(), 9):
                    # Get values with defaults
                    milestone_no = getattr(row, 'milestone_no', f'M{idx-8}') if hasattr(row, 'milestone_no') else f'M{idx-8}'
                    services = getattr(row, 'services', '') if hasattr(row, 'services') else ''
                    due_date = getattr(row, 'due_date', date.today()) if hasattr(row, 'due_date') else date.today()
                    allocation = float(getattr(row, 'allocation', 0)) if hasattr(row, 'allocation') else 0
                    net_pay = float(getattr(row, 'net_pay', 0)) if hasattr(row, 'net_pay') else 0
                    
                    # Write values
                    ws.cell(row=idx, column=1, value=milestone_no)
                    ws.cell(row=idx, column=2, value=services)
                    ws.cell(row=idx, column=3, value=due_date.strftime('%Y-%m-%d') if hasattr(due_date, 'strftime') else str(due_date))
                    ws.cell(row=idx, column=4, value=allocation)
                    ws.cell(row=idx, column=5, value=net_pay)
                    
                    # Add formatting
                    for col in range(1, 6):
                        cell = ws.cell(row=idx, column=col)
                        cell.border = thin_border
                        if col in [4, 5]:
                            cell.number_format = '#,##0.00'
                        if col == 5:
                            cell.font = Font(bold=True)
                
                # Calculate totals
                total_row = len(milestone_df) + 10
                ws.cell(row=total_row, column=4, value="Total:").font = Font(bold=True)
                total_pay = milestone_df['net_pay'].sum() if 'net_pay' in milestone_df.columns else 0
                ws.cell(row=total_row, column=5, value=total_pay)
                ws.cell(row=total_row, column=5).font = Font(bold=True)
                ws.cell(row=total_row, column=5).number_format = '#,##0.00'
            
            # Add summary section
            summary_row = total_row + 2 if milestone_df is not None and not milestone_df.empty else 15
            ws.cell(row=summary_row, column=1, value="Summary").font = Font(bold=True, size=12)
            
            # Save file
            file_name = f"{sow_number}_Milestone_Payments.xlsx"
            file_path = os.path.join(self.output_folder, file_name)
            wb.save(file_path)
            
            print(f"✅ Created milestone Excel: {file_path}")
            return file_path
            
        except Exception as e:
            print(f"❌ Error creating milestone Excel: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def create_tm_resource_excel(self, sow_data, resources_df):
        """Create Excel file for T&M resource details - FIXED VERSION"""
        try:
            # Extract data
            sow_number = sow_data.get("sow_num", "UNKNOWN")
            sow_name = sow_data.get("sow_name", "Unknown SOW")
            client = sow_data.get("Client_Name", "Unknown Client")
            
            # Create workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Resource Details"
            
            # Header styling
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            thin_border = Border(left=Side(style='thin'), 
                                right=Side(style='thin'), 
                                top=Side(style='thin'), 
                                bottom=Side(style='thin'))
            
            # Write SOW Information
            ws['A1'] = f"SOW: {sow_number} - {sow_name}"
            ws['A1'].font = Font(bold=True, size=14, color="4472C4")
            ws.merge_cells('A1:H1')
            
            ws['A2'] = f"Client: {client}"
            ws['A2'].font = Font(bold=True)
            
            ws['A3'] = f"Project Type: T&M (Time & Materials)"
            ws['A3'].font = Font(bold=True)
            
            # Blank row
            ws['A5'] = "Resource Allocation Details"
            ws['A5'].font = Font(bold=True, size=12)
            ws.merge_cells('A5:H5')
            
            # Column headers for resources - FIXED to match your DataFrame
            headers = ["Role", "Location", "Start Date", "End Date", 
                      "Allocation %", "Hrs/Day", "Rate/hr ($)", "Estimated $"]
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=7, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border
                # Adjust column widths
                ws.column_dimensions[get_column_letter(col)].width = 15
            ws.column_dimensions['B'].width = 12  # Location
            ws.column_dimensions['C'].width = 12  # Start Date
            ws.column_dimensions['D'].width = 12  # End Date
            
            # Write resource data
            if resources_df is not None and not resources_df.empty:
                # Make sure we have the right column names
                print(f"🔍 DEBUG: Resources DataFrame columns: {list(resources_df.columns)}")
                print(f"🔍 DEBUG: Resources DataFrame shape: {resources_df.shape}")
                
                for idx, row in resources_df.iterrows():
                    ws_row = idx + 8  # Start from row 8
                    
                    # Safely get values with default fallbacks
                    role = row.get("Role", "")
                    location = row.get("Location", "")
                    start_date = row.get("Start Date", date.today())
                    end_date = row.get("End Date", date.today())
                    allocation = row.get("Allocation %", 0)
                    hrs_day = row.get("Hrs/Day", 8)
                    rate = row.get("Rate/hr ($)", 0)
                    estimated = row.get("Estimated $", 0)
                    
                    # Write values
                    ws.cell(row=ws_row, column=1, value=role)
                    ws.cell(row=ws_row, column=2, value=location)
                    ws.cell(row=ws_row, column=3, value=start_date.strftime('%Y-%m-%d') if hasattr(start_date, 'strftime') else str(start_date))
                    ws.cell(row=ws_row, column=4, value=end_date.strftime('%Y-%m-%d') if hasattr(end_date, 'strftime') else str(end_date))
                    ws.cell(row=ws_row, column=5, value=allocation)
                    ws.cell(row=ws_row, column=6, value=hrs_day)
                    ws.cell(row=ws_row, column=7, value=rate)
                    ws.cell(row=ws_row, column=8, value=estimated)
                    
                    # Add formatting
                    for col in range(1, 9):
                        cell = ws.cell(row=ws_row, column=col)
                        cell.border = thin_border
                        if col in [5, 6, 7, 8]:
                            try:
                                cell.number_format = '#,##0.00'
                            except:
                                pass
                        if col == 8:
                            cell.font = Font(bold=True)
                
                # Calculate totals
                total_row = len(resources_df) + 9
                ws.cell(row=total_row, column=7, value="Total:").font = Font(bold=True)
                
                total_estimated = resources_df["Estimated $"].sum() if "Estimated $" in resources_df.columns else 0
                ws.cell(row=total_row, column=8, value=total_estimated)
                ws.cell(row=total_row, column=8).font = Font(bold=True)
                ws.cell(row=total_row, column=8).number_format = '#,##0.00'
                
                # Add calculation details
                details_row = total_row + 2
                ws.cell(row=details_row, column=1, value="Calculation Method:").font = Font(bold=True)
                ws.cell(row=details_row+1, column=1, 
                       value="Estimated Cost = Working Days × (Allocation%/100) × Hours/Day × Rate/Hour")
                ws.merge_cells(f'A{details_row+1}:H{details_row+1}')
            
            # Add timestamp
            timestamp_row = details_row + 3 if 'details_row' in locals() else 15
            ws.cell(row=timestamp_row, column=1, 
                   value=f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Save file
            file_name = f"{sow_number}_Resource_Details.xlsx"
            file_path = os.path.join(self.output_folder, file_name)
            wb.save(file_path)
            
            print(f"✅ Created resource Excel: {file_path}")
            return file_path
            
        except Exception as e:
            print(f"❌ Error creating resource Excel: {str(e)}")
            import traceback
            traceback.print_exc()
            return None

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================
def init_session_state():
    """Initialize session state variables"""
    defaults = {
        'should_increment_on_download': False,
        'generated_file_path': None,
        'file_data': None,
        'reset_trigger': 0,
        'user_email': "",
        'user_role': 'guest',
        'current_sow_id': None,
        'sharepoint_service': SharePointService(),
        'current_sow_data': None,
        'sow_saved': False,
        'document_uploaded': False,
        'form_data': {},
        'is_authenticated': False,
        # NEW: For approval dashboard persistence
        'sow_dataframe': None,
        'selected_sow_for_download': None,
        'download_triggered': False,
        'document_bytes': None,
        'document_filename': None,
        # NEW: For SOW view/edit mode
        'edit_sow_mode': False,
        'edit_sow_data': None,
        'edit_sow_id': None,
        'edit_sow_status': None,
        'edit_sow_comments': "",
        'viewing_for_approval': False,
        'submit_with_upload': False,  # NEW: Track if submit with upload is checked
        'auto_upload_completed': False,
        'word_document_url': '',  # NEW: Store Word document URL
        'excel_document_url': ''
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def get_base64_image(image_path):
    """Convert image to base64 for HTML display"""
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except:
        return ""

def networkdays(start_date, end_date):
    """Calculate working days (Mon-Fri)"""
    if isinstance(start_date, str):
        start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
    if isinstance(end_date, str):
        end_date = datetime.strptime(end_date, "%Y-%m-%d").date()
    
    # Use numpy for business day calculation
    days = np.busday_count(start_date, end_date)
    return int(days)

def get_next_sow_number():
    """Generate next SOW number"""
    counter_file = "data/sow_counter.txt"
    start_num = 1000
    
    # Create data directory if it doesn't exist
    os.makedirs("data", exist_ok=True)
    
    if not os.path.exists(counter_file):
        with open(counter_file, "w") as f:
            f.write(str(start_num))
        return start_num
    
    try:
        with open(counter_file, "r") as f:
            current = int(f.read().strip())
    except:
        current = start_num
    
    next_num = current + 1
    with open(counter_file, "w") as f:
        f.write(str(next_num))
    
    return current

def reset_all_fields():
    """Reset form fields"""
    keys_to_keep = [
        'user_email', 'user_role', 'sharepoint_service', 
        'sow_saved', 'document_uploaded', 'is_authenticated',
        'edit_sow_mode', 'edit_sow_data', 'edit_sow_id',
        'edit_sow_status', 'edit_sow_comments', 'viewing_for_approval'
    ]
    
    keys_to_remove = [key for key in st.session_state.keys() 
                     if key not in keys_to_keep]
    
    for key in keys_to_remove:
        del st.session_state[key]
    
    st.session_state.reset_trigger = st.session_state.get('reset_trigger', 0) + 1

def save_to_local_csv(sow_data):
    """Save SOW data to local CSV (fallback)"""
    try:
        csv_file = "data/sow_records_local.csv"
        os.makedirs("data", exist_ok=True)
        
        df = pd.DataFrame([sow_data])
        
        if os.path.exists(csv_file):
            existing_df = pd.read_csv(csv_file)
            df = pd.concat([existing_df, df], ignore_index=True)
        
        df.to_csv(csv_file, index=False)
        return True
    except Exception as e:
        st.error(f"Local save failed: {str(e)}")
        return False

def prepare_sow_data_for_storage(form_data, document_url=""):
    """Prepare complete SOW data for SharePoint storage - FIXED"""
    
    def convert_date(obj):
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj
    
    # Calculate total value - FIXED LOGIC
    def calculate_total_value():
        try:
            project_type = form_data.get("option", "")
            
            if project_type == "T&M":
                # Get from form_data if available
                if "currency_value" in form_data:
                    return float(form_data.get("currency_value", 0))
                # Or calculate from resources_df
                resources_df = form_data.get("resources_df")
                if resources_df is not None and not resources_df.empty:
                    if "Estimated $" in resources_df.columns:
                        return float(resources_df["Estimated $"].sum())
                return 0.0
                
            elif project_type == "Fixed Fee":
                # Check both possible keys
                if "Fees_al" in form_data:
                    fees = form_data.get("Fees_al", 0)
                    return float(fees) if fees not in ["", None] else 0.0
                return 0.0
                
            elif project_type == "Change Order":
                # Get difference directly
                diff = form_data.get("difference", 0)
                if diff:
                    return float(diff)
                # Or calculate it
                fees_co = float(form_data.get("Fees_co", 0))
                fees_sow = float(form_data.get("Fees_sow", 0))
                return fees_co - fees_sow
                
            return 0.0
        except Exception as e:
            print(f"❌ Error in calculate_total_value: {e}")
            return 0.0
    
    total_value = calculate_total_value()
    
    # Debug print to check
    print(f"🔍 DEBUG: Total Value Calculation Result:")
    print(f"  - Project Type: {form_data.get('option')}")
    print(f"  - Total Value: {total_value}")
    print(f"  - Form data keys: {list(form_data.keys())}")
    
    # Check specific keys for debugging
    if form_data.get("option") == "Fixed Fee":
        print(f"  - Fees_al in form_data: {'Fees_al' in form_data}")
        if 'Fees_al' in form_data:
            print(f"  - Fees_al value: {form_data['Fees_al']}")
    
    # Calculate work days
    try:
        work_days = networkdays(
            form_data.get("start_date", date.today()),
            form_data.get("end_date", date.today())
        )
    except:
        work_days = 0
    
    # Prepare additional data
    additional_data = {
        "generation_timestamp": datetime.now().isoformat(),
        "template_used": Config.TEMPLATE_MAPPING.get(form_data.get("option", ""), "unknown"),
        "complete_scope": form_data.get("scope_text", ""),
        "complete_services": form_data.get("ser_del", ""),
        "project_specific": {},
        "total_value_debug": total_value
    }
    
    # Add project-specific data
    if form_data.get("option") == "T&M":
        resources_df = form_data.get("resources_df")
        if resources_df is not None and not resources_df.empty:
            additional_data["project_specific"]["resources"] = resources_df.to_dict(orient="records")
            if "Estimated $" in resources_df.columns:
                additional_data["project_specific"]["resources_total"] = resources_df["Estimated $"].sum()
    
    elif form_data.get("option") == "Fixed Fee":
        fees_al = form_data.get("Fees_al", 0)
        additional_data["project_specific"]["fees"] = fees_al
        
        # ✅ ADD THIS: Save milestone data
        milestone_df = form_data.get("milestone_df")
        if milestone_df is not None and not milestone_df.empty:
            # Convert DataFrame to list of dictionaries
            milestones_list = milestone_df.to_dict(orient="records")
            additional_data["project_specific"]["milestones"] = milestones_list
            
            # Calculate and store total
            if "net_pay" in milestone_df.columns:
                total_milestone_payment = milestone_df["net_pay"].sum()
                additional_data["project_specific"]["milestone_total"] = float(total_milestone_payment)
            
            print(f"✅ DEBUG: Saved {len(milestones_list)} milestones to project_specific")
        else:
            print("⚠️ DEBUG: No milestone_df found in form_data for Fixed Fee")
            
            # Save milestone data if available
            milestone_df = form_data.get("milestone_df")
            if milestone_df is not None and not milestone_df.empty:
                # Convert the DataFrame to a list of dictionaries for storage
                milestones_list = milestone_df.to_dict(orient="records")
                additional_data["project_specific"]["milestones"] = milestones_list
                additional_data["project_specific"]["milestone_total"] = milestone_df["net_pay"].sum() if "net_pay" in milestone_df.columns else 0
        
    elif form_data.get("option") == "Change Order":
        additional_data["project_specific"].update({
            "change_order": form_data.get("Change", ""),
            "fees_co": form_data.get("Fees_co", 0),
            "fees_sow": form_data.get("Fees_sow", 0),
            "difference": form_data.get("difference", 0)
        })
    
    # FIX: Ensure we have a valid document URL
    if not document_url or document_url == "":
        document_url = "https://cloudlabsit.sharepoint.com/sites/OnboardingDetails"
    
    # Build SharePoint record - FIXED TotalValue field
    sow_record = {
        "Title": form_data.get("sow_name", ""),
        "SOWNumber": form_data.get("sow_num", ""),
        "SOWName": form_data.get("sow_name", ""),
        "Client": form_data.get("Client_Name", ""),
        "ProjectType": form_data.get("option", ""),
        "Status": Config.STATUS_PENDING,
        "StartDate": convert_date(form_data.get("start_date", date.today())),
        "EndDate": convert_date(form_data.get("end_date", date.today())),
        "GeneratedDate": datetime.now().strftime("%Y-%m-%d"),
        # FIX: Ensure TotalValue is properly set
        "TotalValue": float(total_value) if total_value else 0.0,
        "CreatedBy": st.session_state.user_email,
        "ScopeSummary": form_data.get("scope_text", "")[:1000] if form_data.get("scope_text") else "",
        "ServicesDeliverables": form_data.get("ser_del", "")[:1000] if form_data.get("ser_del") else "",
        "AdditionalPersonnel": form_data.get("additional_personnel", ""),
        "WorkDays": work_days,
        "DocumentURL": document_url,
        "FileName": f"{form_data.get('sow_num', '')} - {form_data.get('sow_name', '')}.docx",
        "PMClient": form_data.get("pm_client", ""),
        "PMServiceProvider": form_data.get("pm_sp", ""),
        "ManagementClient": form_data.get("mg_client", ""),
        "ManagementServiceProvider": form_data.get("mg_sp", ""),
        "AdditionalData": json.dumps(additional_data, default=str)
    }
    
    # Final debug
    print(f"🔍 FINAL SOW RECORD - TotalValue: {sow_record['TotalValue']}")
    print(f"  - Type: {type(sow_record['TotalValue'])}")
    
    return sow_record




    


def collect_form_data_from_session():
    """Collect all form data from session state"""
    form_data = {
        "option": st.session_state.get(f"project_type_{st.session_state.reset_trigger}", ""),
        "sow_num": st.session_state.get(f"sow_num_{st.session_state.reset_trigger}", ""),
        "sow_name": st.session_state.get(f"sow_name_{st.session_state.reset_trigger}", ""),
        "Client_Name": st.session_state.get(f"client_{st.session_state.reset_trigger}", ""),
        "start_date": st.session_state.get(f"start_date_{st.session_state.reset_trigger}", date.today()),
        "end_date": st.session_state.get(f"end_date_{st.session_state.reset_trigger}", date.today()),
        "scope_text": st.session_state.get(f"scope_{st.session_state.reset_trigger}", ""),
        "ser_del": st.session_state.get(f"ser_del_{st.session_state.reset_trigger}", ""),
        "pm_client": st.session_state.get(f"pm_client_{st.session_state.reset_trigger}", ""),
        "pm_sp": st.session_state.get(f"pm_sp_{st.session_state.reset_trigger}", ""),
        "mg_client": st.session_state.get(f"mg_client_{st.session_state.reset_trigger}", ""),
        "mg_sp": st.session_state.get(f"mg_sp_{st.session_state.reset_trigger}", ""),
        "additional_personnel": st.session_state.get(f"additional_personnel_{st.session_state.reset_trigger}", ""),
    }
    
    # Add project-specific data
    option = form_data["option"]
    
    if option == "Fixed Fee":
        form_data["Fees_al"] = st.session_state.get(f"fees_al_{st.session_state.reset_trigger}", 0)
        # Get milestone data from session state
        if hasattr(st.session_state, 'edit_milestone_df'):
            form_data["milestone_df"] = st.session_state.edit_milestone_df
    
    elif option == "T&M":
        # Get resource data from session state
        if hasattr(st.session_state, 'edit_resources_df'):
            form_data["resources_df"] = st.session_state.edit_resources_df
    
    elif option == "Change Order":
        form_data.update({
            "Change": st.session_state.get(f"change_{st.session_state.reset_trigger}", ""),
            "Fees_co": st.session_state.get(f"fees_co_{st.session_state.reset_trigger}", 0),
            "Fees_sow": st.session_state.get(f"fees_sow_{st.session_state.reset_trigger}", 0),
            "difference": st.session_state.get(f"difference_{st.session_state.reset_trigger}", 0)
        })
    
    return form_data



def prepare_sow_data_for_update(form_data, item_id):
    """Prepare SOW data for update"""
    
    # Convert dates to string
    def convert_date(obj):
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj
    
    # Calculate work days
    workdays = networkdays(
        form_data.get("start_date", date.today()),
        form_data.get("end_date", date.today())
    )
    
    # Prepare additional data
    additional_data = {
        "generation_timestamp": datetime.now().isoformat(),
        "last_modified_by": st.session_state.user_email,
        "last_modified_date": datetime.now().isoformat(),
        "complete_scope": form_data.get("scope_text", ""),
        "complete_services": form_data.get("ser_del", ""),
        "project_specific": {}
    }
    
    # Add project-specific data
    if form_data.get("option") == "T&M":
        resources_df = form_data.get("resources_df")
        if resources_df is not None and not resources_df.empty:
            additional_data["project_specific"]["resources"] = resources_df.to_dict(orient="records")
            if "Estimated $" in resources_df.columns:
                additional_data["project_specific"]["resources_total"] = resources_df["Estimated $"].sum()
    
    elif form_data.get("option") == "Fixed Fee":
        additional_data["project_specific"]["fees"] = form_data.get("Fees_al", 0)
        milestone_df = form_data.get("milestone_df")
        if milestone_df is not None and not milestone_df.empty:
            additional_data["project_specific"]["milestones"] = milestone_df.to_dict(orient="records")
            if "net_pay" in milestone_df.columns:
                additional_data["project_specific"]["milestone_total"] = milestone_df["net_pay"].sum()
    
    elif form_data.get("option") == "Change Order":
        additional_data["project_specific"].update({
            "change_order": form_data.get("Change", ""),
            "fees_co": form_data.get("Fees_co", 0),
            "fees_sow": form_data.get("Fees_sow", 0),
            "difference": form_data.get("difference", 0)
        })
    
    # Prepare update payload
    sow_update = {
        "Title": form_data.get("sow_name", ""),
        "SOWName": form_data.get("sow_name", ""),
        "Client": form_data.get("Client_Name", ""),
        "ProjectType": form_data.get("option", ""),
        "StartDate": convert_date(form_data.get("start_date", date.today())),
        "EndDate": convert_date(form_data.get("end_date", date.today())),
        "ScopeSummary": form_data.get("scope_text", "")[:1000],
        "ServicesDeliverables": form_data.get("ser_del", "")[:1000],
        "AdditionalPersonnel": form_data.get("additional_personnel", ""),
        "WorkDays": workdays,
        "PMClient": form_data.get("pm_client", ""),
        "PMServiceProvider": form_data.get("pm_sp", ""),
        "ManagementClient": form_data.get("mg_client", ""),
        "ManagementServiceProvider": form_data.get("mg_sp", ""),
        "AdditionalData": json.dumps(additional_data, default=str)
    }
    
    # Add TotalValue for certain project types
    if form_data.get("option") == "Fixed Fee":
        sow_update["TotalValue"] = float(form_data.get("Fees_al", 0))
    elif form_data.get("option") == "T&M":
        resources_df = form_data.get("resources_df")
        if resources_df is not None and "Estimated $" in resources_df.columns:
            sow_update["TotalValue"] = float(resources_df["Estimated $"].sum())
    elif form_data.get("option") == "Change Order":
        sow_update["TotalValue"] = float(form_data.get("difference", 0))
    
    return sow_update

def save_edited_sow():
    """Save edited SOW data back to SharePoint"""
    with st.spinner("Saving changes..."):
        try:
            # Collect all form data
            form_data = collect_form_data_from_session()
            
            # Prepare update data
            sow_update = prepare_sow_data_for_update(form_data, st.session_state.edit_sow_id)
            
            # Update in SharePoint
            sharepoint_service = st.session_state.sharepoint_service
            result = sharepoint_service.update_sow_record(
                item_id=st.session_state.edit_sow_id,
                sow_data=sow_update
            )
            
            if result["success"]:
                st.success("✅ Changes saved successfully!")
                
                # Update session state with new data
                st.session_state.edit_sow_data.update(sow_update)
                
                time.sleep(2)
                st.rerun()
            else:
                st.error(f"❌ Failed to save changes: {result.get('message', 'Unknown error')}")
                
        except Exception as e:
            st.error(f"❌ Error saving changes: {str(e)}")
# ============================================================================
# LOGIN SYSTEM
# ============================================================================
# ============================================================================
# LOGIN SYSTEM
# ============================================================================
def login_page():
    """Display login page - No password required"""
    st.markdown("""
    <div class="login-container">
        <div class="login-header">
            <h2>📋 SOW Generator</h2>
            <p>Single Click Word SOW Generator | SharePoint Integrated</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Login form
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown("### Login")
    
    with col2:
        email = st.text_input("Email", placeholder="Enter your email")
        
        login_btn = st.button("Login", type="primary", use_container_width=True)
        
        if login_btn:
            if not email:
                st.error("Please enter your email")
            else:
                # Set user session
                st.session_state.is_authenticated = True
                st.session_state.user_email = email
                
                # Determine user role
                if email in Config.LEGAL_TEAM:
                    st.session_state.user_role = 'legal'
                else:
                    st.session_state.user_role = 'user'
                
                st.success(f"Welcome, {email}!")
                time.sleep(1)
                st.rerun()
    
    # Information
    st.markdown("---")
    st.info("**Note:** Just enter your email to access the application. Legal team access is granted only to specific emails.")
    
    st.markdown("</div>", unsafe_allow_html=True)

def logout():
    """Logout user"""
    st.session_state.is_authenticated = False
    st.session_state.user_email = ""
    st.session_state.user_role = 'guest'
    st.rerun()

# ============================================================================
# HEADER COMPONENT
# ============================================================================
def render_header():
    """Render application header with user info and logout button"""
    user_display = st.session_state.user_email.split('@')[0] if '@' in st.session_state.user_email else st.session_state.user_email
    
    st.markdown(f"""
    <div class="custom-header">
        <h1>📋 SOW Generator</h1>
        <p>Single Click Word SOW Generator | SharePoint Integrated</p>
        <div style="position: relative;">
            <button class="logout-btn" onclick="window.parent.document.querySelector('iframe').contentWindow.logout()">
                👤 {user_display} | Logout
            </button>
        </div>
    </div>
    <script>
    function logout() {{
        window.location.href = window.location.href;
    }}
    </script>
    """, unsafe_allow_html=True)

# ============================================================================
# TEMPLATE MANAGEMENT
# ============================================================================
class TemplateManager:
    def __init__(self):
        # Try multiple possible template locations
        self.template_locations = [
            Path("templates"),  # For local development
            Path("sow_app/templates"),  # For Streamlit Cloud if app is in sow_app folder
            Path(".") / "templates",  # Current directory
            Path(__file__).parent / "templates",  # Same directory as main.py
        ]
        
        # Also check for templates in the app directory
        self.ensure_templates_exist()
    
    def ensure_templates_exist(self):
        """Ensure template files exist in the current directory"""
        template_files = [
            "Fixed_Fee_Template.docx",
            "T&M_Template.docx",
            "Change_Order_Template.docx"
        ]
        
        for template_file in template_files:
            found = False
            for location in self.template_locations:
                template_path = location / template_file
                if template_path.exists():
                    print(f"✅ Found template at: {template_path}")
                    found = True
                    break
            
            if not found:
                print(f"⚠️ Template not found: {template_file}")
                # Create a default template
                self.create_default_template_by_name(template_file)
    
    def get_template(self, project_type):
        """Get template based on project type - UPDATED FOR STREAMLIT CLOUD"""
        template_name = Config.TEMPLATE_MAPPING.get(project_type)
        if not template_name:
            st.error(f"No template defined for project type: {project_type}")
            return self.create_default_template(project_type)
        
        # Try to find the template in multiple locations
        for location in self.template_locations:
            template_path = location / template_name
            if template_path.exists():
                try:
                    print(f"✅ Loading template from: {template_path}")
                    return BytesIO(template_path.read_bytes())
                except Exception as e:
                    print(f"❌ Error loading template {template_path}: {e}")
        
        # If not found, create default
        print(f"⚠️ Template {template_name} not found in any location. Creating default.")
        return self.create_default_template(project_type, template_name)
    
    def create_default_template(self, project_type, template_name=None):
        """Create a default template"""
        from docx import Document
        
        doc = Document()
        doc.add_heading(f'Statement of Work - {project_type}', 0)
        
        # Add standard placeholders
        sections = [
            ("SOW Details", [
                f"SOW Number: {{sow_num}}",
                f"SOW Name: {{sow_name}}",
                f"Client: {{client}}",
                f"Project Type: {project_type}"
            ]),
            ("Timeline", [
                "Start Date: {{start_date}}",
                "End Date: {{end_date}}",
                "Generated Date: {{generated_date}}"
            ]),
            ("Scope", ["{{scope_text}}"]),
            ("Services/Deliverables", ["{{ser_del}}"]),
            ("Financial Details", [
                "{% if project_type == 'Fixed Fee' %}",
                "Total Fees: {{Fees}}",
                "{% elif project_type == 'T&M' %}",
                "Total Value: {{currency_value_str}}",
                "{% elif project_type == 'Change Order' %}",
                "Change Order: {{Change}}",
                "Fees CO: {{Fees_co}}",
                "Fees SOW: {{Fees_sow}}",
                "Difference: {{difference}}",
                "{% endif %}"
            ]),
            ("Personnel", ["{{additional_personnel}}"]),
            ("Contacts", [
                "PM Client: {{pm_client}}",
                "PM Service Provider: {{pm_sp}}",
                "Management Client: {{mg_client}}",
                "Management Service Provider: {{mg_sp}}"
            ])
        ]
        
        for section_title, section_content in sections:
            doc.add_heading(section_title, level=1)
            for line in section_content:
                doc.add_paragraph(line)
            doc.add_paragraph()  # Empty line
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Save locally for future use
        if template_name:
            # Try to save in the first available location
            for location in self.template_locations:
                try:
                    location.mkdir(exist_ok=True)
                    save_path = location / template_name
                    save_path.write_bytes(buffer.getvalue())
                    print(f"💾 Saved template to: {save_path}")
                    break
                except Exception as e:
                    print(f"Could not save to {location}: {e}")
        
        return buffer
    
    def create_default_template_by_name(self, template_name):
        """Create a default template based on template name"""
        # Determine project type from template name
        if "Fixed_Fee" in template_name:
            project_type = "Fixed Fee"
        elif "T&M" in template_name:
            project_type = "T&M"
        elif "Change_Order" in template_name:
            project_type = "Change Order"
        else:
            project_type = "Fixed Fee"
        
        return self.create_default_template(project_type, template_name)



def load_sow_data_for_edit_mode(sow_data):
    """Load SOW data into session state for edit mode with editable dataframes"""
    try:
        # Parse additional data
        additional_data = {}
        try:
            additional_data = json.loads(sow_data.get("AdditionalData", "{}"))
        except:
            additional_data = {}
        
        # Extract resources for T&M projects - make editable
        if sow_data.get("ProjectType") == "T&M":
            resources_list = additional_data.get("project_specific", {}).get("resources", [])
            if resources_list:
                resources_df = pd.DataFrame(resources_list)
                st.session_state.edit_resources_df = resources_df
                
                # Store as editable dataframe in session state
                st.session_state.editable_resources = resources_df.copy()
        
        # Extract milestones for Fixed Fee projects - make editable
        elif sow_data.get("ProjectType") == "Fixed Fee":
            milestones_list = additional_data.get("project_specific", {}).get("milestones", [])
            if milestones_list:
                milestone_df = pd.DataFrame(milestones_list)
                st.session_state.edit_milestone_df = milestone_df
                
                # Store as editable dataframe in session state
                st.session_state.editable_milestones = milestone_df.copy()
        
        return True
    except Exception as e:
        print(f"Error loading SOW data for edit mode: {str(e)}")
        return False
# ============================================================================
# PAGE 1: SOW GENERATOR (UPDATED WITH EDIT MODE)
# ============================================================================
def page_sow_generator():
    """Main SOW generation page - Updated with edit mode"""
    
    # Check if we're in edit mode (viewing for approval)
    if st.session_state.edit_sow_mode:
        st.markdown("<div class='approval-section'>", unsafe_allow_html=True)
        st.title("👁️ View SOW for Approval")
        
        # Display SOW info
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("SOW Number", st.session_state.edit_sow_data.get("SOWNumber", ""))
        with col2:
            st.metric("Status", st.session_state.edit_sow_data.get("Status", ""))
        with col3:
            st.metric("Created By", st.session_state.edit_sow_data.get("CreatedBy", ""))
        
        st.divider()
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    template_manager = TemplateManager()
    
    # Show template info
    if not st.session_state.edit_sow_mode:
        st.info("📝 **Templates are automatically selected based on project type**")
    
    # ========== BASIC INFORMATION ==========
    st.subheader("📋 Basic Information")
    col1, col2 = st.columns(2)
    
    with col1:
        # If in edit mode, use existing data, otherwise create new form
        if st.session_state.edit_sow_mode:
            Client_Name = st.selectbox(
                "Select Client",
                ("BSC", "Abiomed", "Cognex", "Itaros", "Other"),
                key=f"client_{st.session_state.reset_trigger}",
                help="Select the client for this SOW",
                index=["BSC", "Abiomed", "Cognex", "Itaros", "Other"].index(
                    st.session_state.edit_sow_data.get("Client", "BSC")
                ) if st.session_state.edit_sow_data.get("Client") in ["BSC", "Abiomed", "Cognex", "Itaros", "Other"] else 0,
                disabled= True
            )
            
            option = st.selectbox(
                "Select Project Type",
                ("Fixed Fee", "T&M", "Change Order"),
                key=f"project_type_{st.session_state.reset_trigger}",
                help="Select the project type to automatically choose the template",
                index=["Fixed Fee", "T&M", "Change Order"].index(
                    st.session_state.edit_sow_data.get("ProjectType", "Fixed Fee")
                ) if st.session_state.edit_sow_data.get("ProjectType") in ["Fixed Fee", "T&M", "Change Order"] else 0,
                disabled= True
            )
        else:
            Client_Name = st.selectbox(
                "Select Client",
                ("BSC", "Abiomed", "Cognex", "Itaros", "Other"),
                key=f"client_{st.session_state.reset_trigger}",
                help="Select the client for this SOW"
            )
            
            option = st.selectbox(
                "Select Project Type",
                ("Fixed Fee", "T&M", "Change Order"),
                key=f"project_type_{st.session_state.reset_trigger}",
                help="Select the project type to automatically choose the template"
            )
        
        # Show which template will be used
        template_name = Config.TEMPLATE_MAPPING.get(option, "Unknown")
        st.caption(f"Template: **{template_name}**")
    
    with col2:
        # SOW Number generation
        if st.session_state.edit_sow_mode:
            sow_num = st.text_input(
                "SOW Number",
                value=st.session_state.edit_sow_data.get("SOWNumber", ""),
                key=f"sow_num_{st.session_state.reset_trigger}",
                help="Auto-generated SOW number",
                disabled=True  # Disable editing of SOW number in view mode
            )
            sow_name = st.text_input(
                "SOW Name",
                value=st.session_state.edit_sow_data.get("SOWName", ""),
                key=f"sow_name_{st.session_state.reset_trigger}",
                placeholder="Enter SOW name",
                help="Enter a descriptive name for this SOW",
                disabled=True  # Disable editing of SOW name in view mode
            )
        else:
            if option in ["T&M", "Fixed Fee"]:
                auto_sow_num = get_next_sow_number()
                sow_num = st.text_input(
                    "SOW Number",
                    value=f"SOW-{auto_sow_num}",
                    key=f"sow_num_{st.session_state.reset_trigger}",
                    help="Auto-generated SOW number"
                )
            else:
                sow_num = st.text_input(
                    "SOW Number",
                    value="",
                    key=f"sow_num_{st.session_state.reset_trigger}",
                    placeholder="Enter SOW number manually for Change Order",
                    help="Enter SOW number for Change Order"
                )
            
            sow_name = st.text_input(
                "SOW Name",
                key=f"sow_name_{st.session_state.reset_trigger}",
                placeholder="Enter SOW name",
                help="Enter a descriptive name for this SOW"
            )
    
    # ========== DATES ==========
    st.subheader("📅 Timeline")
    colA, colB = st.columns(2)
    
    with colA:
        if st.session_state.edit_sow_mode:
            # Parse date from stored data
            start_date_str = st.session_state.edit_sow_data.get("StartDate", date.today().isoformat())
            try:
                start_date = datetime.fromisoformat(start_date_str).date()
            except:
                start_date = date.today()
            
            start_date = st.date_input(
                "Start Date",
                value=start_date,
                key=f"start_date_{st.session_state.reset_trigger}",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
        else:
            start_date = st.date_input(
                "Start Date",
                value=date.today(),
                key=f"start_date_{st.session_state.reset_trigger}"
            )
    
    with colB:
        if st.session_state.edit_sow_mode:
            # Parse date from stored data
            end_date_str = st.session_state.edit_sow_data.get("EndDate", (date.today() + timedelta(days=30)).isoformat())
            try:
                end_date = datetime.fromisoformat(end_date_str).date()
            except:
                end_date = date.today() + timedelta(days=30)
            
            end_date = st.date_input(
                "End Date",
                value=end_date,
                key=f"end_date_{st.session_state.reset_trigger}",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
        else:
            end_date = st.date_input(
                "End Date",
                value=date.today() + timedelta(days=30),
                key=f"end_date_{st.session_state.reset_trigger}"
            )
    
    # Calculate and display working days
    workdays = networkdays(start_date, end_date)
    st.info(f"📅 **Total working days (Mon–Fri): {workdays} days**")
    
    # ========== CHANGE ORDER SPECIFIC ==========
    if option == "Change Order":
        st.subheader("🔄 Change Order Details")
        colA, colB = st.columns(2)
        
        # Parse additional data if in edit mode
        additional_data = {}
        if st.session_state.edit_sow_mode:
            try:
                additional_data = json.loads(st.session_state.edit_sow_data.get("AdditionalData", "{}"))
            except:
                additional_data = {}
        
        with colA:
            if st.session_state.edit_sow_mode:
                Change = st.text_input(
                    "Change Order Reference",
                    value=additional_data.get("project_specific", {}).get("change_order", "CO-001"),
                    key=f"change_{st.session_state.reset_trigger}",
                    help="Change order reference number",
                    disabled=not st.session_state.get('edit_mode_enabled', False)
                )
                
                sow_start_date = st.date_input(
                    "Original SOW Start Date",
                    value=date.today(),
                    key=f"sow_start_{st.session_state.reset_trigger}",
                    disabled=not st.session_state.get('edit_mode_enabled', False)
                )
                
                Fees_co = st.number_input(
                    "Change Order Fees ($)",
                    value=float(additional_data.get("project_specific", {}).get("fees_co", 10000.0)),
                    step=1000.0,
                    key=f"fees_co_{st.session_state.reset_trigger}",
                    disabled=not st.session_state.get('edit_mode_enabled', False)
                )
            else:
                Change = st.text_input(
                    "Change Order Reference",
                    value="CO-001",
                    key=f"change_{st.session_state.reset_trigger}",
                    help="Change order reference number"
                )
                
                sow_start_date = st.date_input(
                    "Original SOW Start Date",
                    value=date.today(),
                    key=f"sow_start_{st.session_state.reset_trigger}"
                )
                
                Fees_co = st.number_input(
                    "Change Order Fees ($)",
                    value=10000.0,
                    step=1000.0,
                    key=f"fees_co_{st.session_state.reset_trigger}"
                )
        
        with colB:
            if st.session_state.edit_sow_mode:
                sow_end_date = st.date_input(
                    "Original SOW End Date",
                    value=date.today() + timedelta(days=30),
                    key=f"sow_end_{st.session_state.reset_trigger}",
                    disabled=not st.session_state.get('edit_mode_enabled', False)
                )
                
                Fees_sow = st.number_input(
                    "Original SOW Fees ($)",
                    value=float(additional_data.get("project_specific", {}).get("fees_sow", 5000.0)),
                    step=1000.0,
                    key=f"fees_sow_{st.session_state.reset_trigger}",
                    disabled=not st.session_state.get('edit_mode_enabled', False)
                )
            else:
                sow_end_date = st.date_input(
                    "Original SOW End Date",
                    value=date.today() + timedelta(days=30),
                    key=f"sow_end_{st.session_state.reset_trigger}"
                )
                
                Fees_sow = st.number_input(
                    "Original SOW Fees ($)",
                    value=5000.0,
                    step=1000.0,
                    key=f"fees_sow_{st.session_state.reset_trigger}"
                )
        
        difference = Fees_co - Fees_sow
        st.info(f"💰 **Difference: ${difference:,.2f}**")
    
    # ========== CLIENT & PROVIDER DETAILS ==========
    st.subheader("👥 Client & Service Provider Details")
    colA, colB = st.columns(2)
    
    with colA:
        if st.session_state.edit_sow_mode:
            pm_client = st.text_input(
                "Client (Project Management)",
                value=st.session_state.edit_sow_data.get("PMClient", ""),
                key=f"pm_client_{st.session_state.reset_trigger}",
                placeholder="Client project manager name",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
            
            mg_client = st.text_input(
                "Client (Management)",
                value=st.session_state.edit_sow_data.get("ManagementClient", ""),
                key=f"mg_client_{st.session_state.reset_trigger}",
                placeholder="Client management contact",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
        else:
            pm_client = st.text_input(
                "Client (Project Management)",
                key=f"pm_client_{st.session_state.reset_trigger}",
                placeholder="Client project manager name"
            )
            
            mg_client = st.text_input(
                "Client (Management)",
                key=f"mg_client_{st.session_state.reset_trigger}",
                placeholder="Client management contact"
            )
    
    with colB:
        if st.session_state.edit_sow_mode:
            pm_sp = st.text_input(
                "Service Provider (Project Management)",
                value=st.session_state.edit_sow_data.get("PMServiceProvider", ""),
                key=f"pm_sp_{st.session_state.reset_trigger}",
                placeholder="Service provider project manager",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
            
            mg_sp = st.text_input(
                "Service Provider (Management)",
                value=st.session_state.edit_sow_data.get("ManagementServiceProvider", ""),
                key=f"mg_sp_{st.session_state.reset_trigger}",
                placeholder="Service provider management contact",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
        else:
            pm_sp = st.text_input(
                "Service Provider (Project Management)",
                key=f"pm_sp_{st.session_state.reset_trigger}",
                placeholder="Service provider project manager"
            )
            
            mg_sp = st.text_input(
                "Service Provider (Management)",
                key=f"mg_sp_{st.session_state.reset_trigger}",
                placeholder="Service provider management contact"
            )
    
    # ========== SCOPE & SERVICES ==========
    st.subheader("🎯 Scope & Services")
    
    if st.session_state.edit_sow_mode:
        scope_text = st.text_area(
            "Scope / Responsibilities",
            value=st.session_state.edit_sow_data.get("ScopeSummary", ""),
            height=150,
            key=f"scope_{st.session_state.reset_trigger}",
            placeholder="Describe the scope and responsibilities...",
            disabled=not st.session_state.get('edit_mode_enabled', False)
        )
        
        ser_del = st.text_area(
            "Services / Deliverables",
            value=st.session_state.edit_sow_data.get("ServicesDeliverables", ""),
            height=150,
            key=f"ser_del_{st.session_state.reset_trigger}",
            placeholder="List services and deliverables...",
            disabled=not st.session_state.get('edit_mode_enabled', False)
        )
    else:
        scope_text = st.text_area(
            "Scope / Responsibilities",
            height=150,
            key=f"scope_{st.session_state.reset_trigger}",
            placeholder="Describe the scope and responsibilities..."
        )
        
        ser_del = st.text_area(
            "Services / Deliverables",
            height=150,
            key=f"ser_del_{st.session_state.reset_trigger}",
            placeholder="List services and deliverables..."
        )
    
    # ========== FINANCIAL DETAILS ==========
    if option == "Fixed Fee":
        st.subheader("💰 Fixed Fee Details")
        
        if st.session_state.edit_sow_mode:
            # Parse additional data
            additional_data = {}
            try:
                additional_data = json.loads(st.session_state.edit_sow_data.get("AdditionalData", "{}"))
            except:
                additional_data = {}
            
            Fees_al = st.number_input(
                "Total Fees ($)",
                value=float(additional_data.get("project_specific", {}).get("fees", 0)),
                step=1000.0,
                key=f"fees_al_{st.session_state.reset_trigger}",
                disabled=not st.session_state.get('edit_mode_enabled', False)
            )
        else:
            Fees_al = st.number_input(
                "Total Fees ($)",
                value=50000.0,
                step=1000.0,
                key=f"fees_al_{st.session_state.reset_trigger}"
            )
    
    # ========== ADDITIONAL PERSONNEL ==========
    if st.session_state.edit_sow_mode:
        additional_personnel = st.text_input(
            "Additional Personnel",
            value=st.session_state.edit_sow_data.get("AdditionalPersonnel", ""),
            key=f"additional_personnel_{st.session_state.reset_trigger}",
            placeholder="List any additional personnel involved...",
            disabled=not st.session_state.get('edit_mode_enabled', False)
        )
    else:
        additional_personnel = st.text_input(
            "Additional Personnel",
            key=f"additional_personnel_{st.session_state.reset_trigger}",
            placeholder="List any additional personnel involved..."
        )
    
    # ========== T&M RESOURCES TABLE ==========
    # ========== T&M RESOURCES TABLE ==========
    # ========== T&M RESOURCES TABLE ==========
   # ========== T&M RESOURCES TABLE ==========
    resources_df = None
    if option == "T&M":
        st.subheader("👥 Resource Details")
        
        # Parse resources from additional data if in edit mode
        if st.session_state.edit_sow_mode:
            additional_data = {}
            try:
                additional_data = json.loads(st.session_state.edit_sow_data.get("AdditionalData", "{}"))
            except:
                additional_data = {}
            
            resources_list = additional_data.get("project_specific", {}).get("resources", [])
            
            if resources_list:
                # Convert to DataFrame
                resources_df = pd.DataFrame(resources_list)
                
                # Convert date strings to date objects for editing
                for date_col in ['Start Date', 'End Date']:
                    if date_col in resources_df.columns:
                        resources_df[date_col] = pd.to_datetime(resources_df[date_col]).dt.date
                
                # Make it editable if in edit mode
                if st.session_state.get('edit_mode_enabled', False):
                    st.subheader("✏️ Edit Resource Details")
                    
                    edited_df = st.data_editor(
                        resources_df,
                        num_rows="dynamic",
                        column_config={
                            "Role": st.column_config.TextColumn("Role", width="medium", required=True),
                            "Location": st.column_config.TextColumn("Location", width="small", required=True),
                            "Start Date": st.column_config.DateColumn("Start Date", format="YYYY-MM-DD", required=True),
                            "End Date": st.column_config.DateColumn("End Date", format="YYYY-MM-DD", required=True),
                            "Allocation %": st.column_config.NumberColumn("Allocation %", min_value=0, max_value=100, step=5, required=True, format="%d%%"),
                            "Hrs/Day": st.column_config.NumberColumn("Hrs/Day", min_value=1, max_value=24, step=1, required=True),
                            "Rate/hr ($)": st.column_config.NumberColumn("Rate/hr ($)", min_value=0, step=10, required=True, format="$%.2f"),
                            "Estimated $": st.column_config.NumberColumn("Estimated $", format="$%.2f", disabled=True)
                        },
                        key=f"edit_resources_{st.session_state.reset_trigger}",
                        hide_index=True,
                        use_container_width=True
                    )
                    
                    # Recalculate Estimated $ if needed
                    if not edited_df.empty:
                        def calculate_cost(row):
                            try:
                                days = networkdays(row["Start Date"], row["End Date"])
                                return round(days * (row["Allocation %"]/100) * row["Hrs/Day"] * row["Rate/hr ($)"], 2)
                            except:
                                return 0.0
                        
                        edited_df["Estimated $"] = edited_df.apply(calculate_cost, axis=1)
                    
                    # Store in session state
                    st.session_state.editable_resources = edited_df
                    st.session_state.edit_resources_df = edited_df
                    resources_df = edited_df
                    
                    # Calculate total for display
                    if "Estimated $" in edited_df.columns:
                        total_estimated = edited_df["Estimated $"].sum()
                        st.success(f"💰 **Total Estimated Cost: ${total_estimated:,.2f}**")
                else:
                    # Display as read-only
                    st.dataframe(
                        resources_df,
                        use_container_width=True,
                        height=300
                    )
                    
                    # Calculate total for display
                    if "Estimated $" in resources_df.columns:
                        total_estimated = resources_df["Estimated $"].sum()
                        st.success(f"💰 **Total Estimated Cost: ${total_estimated:,.2f}**")
            else:
                st.info("No resource details available")
        else:
            # Create default resource data for new SOW
            default_data = [{
                "Role": "Senior Consultant",
                "Location": "Remote",
                "Start Date": start_date,
                "End Date": end_date,
                "Allocation %": 100,
                "Hrs/Day": 8,
                "Rate/hr ($)": 150
            }]
            
            # Create editable table
            resources_df = st.data_editor(
                pd.DataFrame(default_data),
                num_rows="dynamic",
                column_config={
                    "Role": st.column_config.TextColumn("Role", width="medium", required=True),
                    "Location": st.column_config.TextColumn("Location", width="small", required=True),
                    "Start Date": st.column_config.DateColumn("Start Date", format="YYYY-MM-DD", required=True),
                    "End Date": st.column_config.DateColumn("End Date", format="YYYY-MM-DD", required=True),
                    "Allocation %": st.column_config.NumberColumn("Allocation %", min_value=0, max_value=100, step=5, required=True, format="%d%%"),
                    "Hrs/Day": st.column_config.NumberColumn("Hrs/Day", min_value=1, max_value=24, step=1, required=True),
                    "Rate/hr ($)": st.column_config.NumberColumn("Rate/hr ($)", min_value=0, step=10, required=True, format="$%.2f")
                },
                key=f"resources_table_{st.session_state.reset_trigger}",
                hide_index=True
            )
            
            # Calculate values
            if not resources_df.empty:
                def calculate_resource_cost(row):
                    try:
                        days = networkdays(row["Start Date"], row["End Date"])
                        return round(days * (row["Allocation %"]/100) * row["Hrs/Day"] * row["Rate/hr ($)"], 2)
                    except:
                        return 0.0
                
                resources_df["Estimated $"] = resources_df.apply(calculate_resource_cost, axis=1)
                
                # Display calculated table
                st.dataframe(resources_df)
                
                # Calculate total
                currency_value = resources_df["Estimated $"].sum()
                currency_value_str = f"${currency_value:,.2f}"
                st.success(f"💰 **Total Contract Value: {currency_value_str}**")
    
    # ========== FIXED FEE MILESTONES ==========
    # ========== FIXED FEE MILESTONES ==========
    # ========== FIXED FEE MILESTONES ==========
    # ========== FIXED FEE MILESTONES ==========
    # ========== FIXED FEE MILESTONES ==========
    milestone_df = None
    if option == "Fixed Fee":
        st.subheader("📊 Milestone Schedule / Payment Breakdown")
        
        if st.session_state.edit_sow_mode:
            # Try to parse milestones from additional data
            additional_data = {}
            try:
                additional_data = json.loads(st.session_state.edit_sow_data.get("AdditionalData", "{}"))
            except:
                additional_data = {}
            
            milestones_list = additional_data.get("project_specific", {}).get("milestones", [])
            
            if milestones_list:
                # Convert to DataFrame
                milestone_df = pd.DataFrame(milestones_list)
                
                # Convert date strings to date objects for editing
                if 'due_date' in milestone_df.columns:
                    # Convert string dates to date objects
                    milestone_df['due_date'] = pd.to_datetime(milestone_df['due_date']).dt.date
                
                # Make it editable if in edit mode
                if st.session_state.get('edit_mode_enabled', False):
                    st.subheader("✏️ Edit Milestone Schedule")
                    
                    edited_df = st.data_editor(
                        milestone_df,
                        num_rows="dynamic",
                        column_config={
                            "milestone_no": st.column_config.TextColumn("Milestone #", width="small"),
                            "services": st.column_config.TextColumn("Services / Deliverables", width="large"),
                            "due_date": st.column_config.DateColumn("Due Date", format="YYYY-MM-DD", required=True),
                            "allocation": st.column_config.NumberColumn("Allocation %", min_value=0, max_value=100, step=5, required=True, format="%d%%"),
                            "net_pay": st.column_config.NumberColumn("Payment ($)", format="$%.2f", disabled=True)
                        },
                        key=f"edit_milestones_{st.session_state.reset_trigger}",
                        hide_index=True,
                        use_container_width=True
                    )
                    
                    # Get current fees from session state
                    current_fees = st.session_state.get(f"fees_al_{st.session_state.reset_trigger}", 0)
                    
                    # Recalculate net_pay if needed
                    if not edited_df.empty and "allocation" in edited_df.columns:
                        edited_df["net_pay"] = edited_df["allocation"].apply(lambda x: round(current_fees * (x/100), 2))
                    
                    # Store in session state
                    st.session_state.editable_milestones = edited_df
                    st.session_state.edit_milestone_df = edited_df
                    milestone_df = edited_df
                    
                    # Calculate total for display
                    if "net_pay" in edited_df.columns:
                        total_payment = edited_df["net_pay"].sum()
                        st.success(f"💰 **Total Milestone Payment: ${total_payment:,.2f}**")
                else:
                    # Display as read-only
                    st.dataframe(
                        milestone_df,
                        use_container_width=True,
                        height=300
                    )
                    
                    # Calculate and display total
                    if "net_pay" in milestone_df.columns:
                        total_payment = milestone_df["net_pay"].sum()
                        st.success(f"💰 **Total Milestone Payment: ${total_payment:,.2f}**")
            else:
                st.info("No milestone details available")
        else:
            # Create default milestone data for new SOW
            default_data = [{
                "Milestone #": "1",
                "Services / Deliverables": "Project Kickoff and Requirements Gathering",
                "Milestone Due Date": start_date,
                "Payment Allocation (%)": 30
            }]
            
            # Create editable table
            milestone_input_df = st.data_editor(
                pd.DataFrame(default_data),
                num_rows="dynamic",
                column_config={
                    "Milestone #": st.column_config.TextColumn("Milestone #", width="small"),
                    "Services / Deliverables": st.column_config.TextColumn("Services / Deliverables", width="large"),
                    "Milestone Due Date": st.column_config.DateColumn("Due Date", format="YYYY-MM-DD", required=True),
                    "Payment Allocation (%)": st.column_config.NumberColumn("Allocation %", min_value=0, max_value=100, step=5, required=True, format="%d%%")
                },
                key=f"milestone_table_{st.session_state.reset_trigger}",
                hide_index=True
            )
            
            # Calculate payments
            try:
                total_fees = float(Fees_al) if 'Fees_al' in locals() else 0
            except:
                total_fees = 0
            
            milestone_df = milestone_input_df.copy()
            
            def calculate_payment(row):
                try:
                    alloc = float(row["Payment Allocation (%)"])
                    return round(total_fees * (alloc / 100), 2)
                except:
                    return 0
            
            milestone_df["Net Milestone Payment ($)"] = milestone_df.apply(calculate_payment, axis=1)
            
            # Display calculated table
            st.dataframe(milestone_df)
            
            # Calculate total
            total_payment = milestone_df["Net Milestone Payment ($)"].sum()
            st.success(f"✅ **Total Net Milestone Payment: ${total_payment:,.2f}**")
            
            # Format for template
            milestone_df = milestone_df.rename(columns={
                "Milestone #": "milestone_no",
                "Services / Deliverables": "services",
                "Milestone Due Date": "due_date",
                "Payment Allocation (%)": "allocation",
                "Net Milestone Payment ($)": "net_pay"
            })
    # ========== APPROVAL SECTION (ONLY IN EDIT MODE) ==========
    if st.session_state.edit_sow_mode and st.session_state.viewing_for_approval:
        st.markdown("</div>", unsafe_allow_html=True)  # Close approval section div
        st.divider()
        
        # Show current status
        current_status = st.session_state.edit_sow_data.get("Status", "")
        status_color = {
            Config.STATUS_PENDING: "🟡",
            Config.STATUS_APPROVED: "🟢",
            Config.STATUS_REJECTED: "🔴",
            Config.STATUS_DRAFT: "⚪"
        }.get(current_status, "⚪")
        
        st.subheader(f"{status_color} Approval Decision")
        
        # Approval decision
        col1, col2 = st.columns(2)
        
        with col1:
            approve_btn = st.button("✅ Approve", use_container_width=True, 
                                   type="primary", key="approve_btn")
        
        with col2:
            reject_btn = st.button("❌ Reject", use_container_width=True,
                                  type="secondary", key="reject_btn")
        
        # Comments
        st.text_area("Comments (Required for Rejection)", 
                    key="approval_comments",
                    placeholder="Enter your comments here...",
                    height=100)
        
        # Handle approval/rejection
        if approve_btn:
            if handle_approval_rejection(Config.STATUS_APPROVED):
                st.success("SOW Approved successfully!")
                time.sleep(2)
                # Reset edit mode and go back to approval dashboard
                st.session_state.edit_sow_mode = False
                st.session_state.viewing_for_approval = False
                st.rerun()
        
        if reject_btn:
            if not st.session_state.get("approval_comments", "").strip():
                st.error("Please provide comments when rejecting a SOW.")
            else:
                if handle_approval_rejection(Config.STATUS_REJECTED):
                    st.success("SOW Rejected successfully!")
                    time.sleep(2)
                    # Reset edit mode and go back to approval dashboard
                    st.session_state.edit_sow_mode = False
                    st.session_state.viewing_for_approval = False
                    st.rerun()

        # ========== SAVE CHANGES BUTTON (EDIT MODE ONLY) ==========
        if st.session_state.edit_sow_mode and st.session_state.get('edit_mode_enabled', False):
            st.divider()
            col1, col2, col3 = st.columns([1, 2, 1])
            
            with col2:
                if st.button("💾 Save Changes", type="primary", use_container_width=True):
                    save_edited_sow()
        
        # Back button
        if st.button("⬅️ Back to Approval Dashboard", use_container_width=True):
            st.session_state.edit_sow_mode = False
            st.session_state.viewing_for_approval = False
            st.rerun()
    
    # ========== GENERATE BUTTON (ONLY IN CREATE MODE) ==========
    # ========== GENERATE BUTTON (ONLY IN CREATE MODE) ==========
    elif not st.session_state.edit_sow_mode:
        st.divider()
        
        # ADD THIS CHECKBOX BEFORE THE GENERATE BUTTON
        col_checkbox, col_info = st.columns([1, 3])
        with col_checkbox:
            submit_with_upload = st.checkbox(
                "✅ Generate SOW with Submit request (Auto-upload to SharePoint)", 
                value=False,
                key=f"submit_with_upload_{st.session_state.reset_trigger}",
                help="Check this box to generate SOW and automatically upload to SharePoint"
            )
            
            # Store checkbox state in session state
            st.session_state.submit_with_upload = submit_with_upload

        with col_info:
            if submit_with_upload:
                st.info("🔄 When you click 'Submit SOW Request', the document will be automatically uploaded to SharePoint after generation.")
            else:
                st.info("📝 When you click 'Submit SOW Request', you'll need to manually upload to SharePoint using the upload button.")

        col1, col2, col3 = st.columns([1, 2, 1])

        with col2:
            generate_btn = st.button(
                "🚀 Submit SOW Request",
                type="primary",
                use_container_width=True,
                help="Click to generate the SOW document"
            )

        if generate_btn:
            generate_sow_document(
                option=option,
                sow_num=sow_num,
                sow_name=sow_name,
                Client_Name=Client_Name,
                start_date=start_date,
                end_date=end_date,
                scope_text=scope_text,
                ser_del=ser_del,
                pm_client=pm_client,
                pm_sp=pm_sp,
                mg_client=mg_client,
                mg_sp=mg_sp,
                additional_personnel=additional_personnel,
                resources_df=resources_df,
                milestone_df=milestone_df,
                template_manager=template_manager,
                auto_upload=submit_with_upload  # Pass the checkbox value
            )
    
    # ========== DOWNLOAD SECTION ==========
    if st.session_state.get('should_increment_on_download') and st.session_state.get('file_data'):
        show_download_section()

def generate_sow_document(option, sow_num, sow_name, Client_Name, start_date, end_date,
                         scope_text, ser_del, pm_client, pm_sp, mg_client, mg_sp,
                         additional_personnel, resources_df, milestone_df, template_manager,
                         auto_upload=False):  # NEW: Add auto_upload parameter
    """Generate SOW document - FIXED VERSION"""
    
    with st.spinner("Generating SOW document..."):
        try:
            # Format dates
            generated_date = datetime.today().strftime("%B %d, %Y")
            start_str = start_date.strftime("%B %d, %Y")
            end_str = end_date.strftime("%B %d, %Y")
            
            # Get financial values BEFORE creating context
            # For T&M projects
            currency_value = 0.0
            if option == "T&M" and resources_df is not None:
                if "Estimated $" in resources_df.columns:
                    currency_value = resources_df["Estimated $"].sum()
            
            # For Fixed Fee projects - IMPORTANT: Get the Fees_al value from session state
            Fees_al_value = 0.0
            if option == "Fixed Fee":
                # Try to get the Fees_al value from the current form
                Fees_al_value = st.session_state.get(f"fees_al_{st.session_state.reset_trigger}", 0.0)
            
            # For Change Order projects
            difference_value = 0.0
            Fees_co_value = 0.0
            Fees_sow_value = 0.0
            if option == "Change Order":
                # Try to get values from session state
                Fees_co_value = st.session_state.get(f"fees_co_{st.session_state.reset_trigger}", 10000.0)
                Fees_sow_value = st.session_state.get(f"fees_sow_{st.session_state.reset_trigger}", 5000.0)
                difference_value = Fees_co_value - Fees_sow_value
            
            # Prepare context based on project type
            context = {
                "sow_num": sow_num,
                "sow_name": sow_name,
                "client": Client_Name,
                "project_type": option,
                "start_date": start_str,
                "end_date": end_str,
                "generated_date": generated_date,
                "scope_text": scope_text,
                "ser_del": ser_del,
                "pm_client": pm_client,
                "pm_sp": pm_sp,
                "mg_client": mg_client,
                "mg_sp": mg_sp,
                "additional_personnel": additional_personnel
            }
            
            # Add project-specific data
            if option == "T&M" and resources_df is not None:
                context["resources"] = resources_df.to_dict(orient="records")
                context["currency_value"] = currency_value
                context["currency_value_str"] = f"${currency_value:,.2f}"
                
                # Store resources in session state for later use
                if st.session_state.edit_sow_mode:
                    st.session_state.form_data["resources_df"] = resources_df
            
            elif option == "Fixed Fee":
                context["Fees"] = Fees_al_value
                if milestone_df is not None:
                    context["milestones"] = milestone_df.to_dict(orient="records")
                    total_payment = milestone_df["net_pay"].sum() if "net_pay" in milestone_df.columns else 0
                    context["milestone_total"] = total_payment
                    
                    # Store milestones in session state for later use
                    if st.session_state.edit_sow_mode:
                        st.session_state.form_data["milestone_df"] = milestone_df
            
            elif option == "Change Order":
                context.update({
                    "Change": st.session_state.get(f"change_{st.session_state.reset_trigger}", "CO-001"),
                    "Fees_co": Fees_co_value,
                    "Fees_sow": Fees_sow_value,
                    "difference": difference_value,
                    "sow_str": st.session_state.get(f"sow_start_{st.session_state.reset_trigger}", date.today()).strftime("%B %d, %Y"),
                    "sow_end": st.session_state.get(f"sow_end_{st.session_state.reset_trigger}", date.today()).strftime("%B %d, %Y")
                })
            
            # Get template
            template_stream = template_manager.get_template(option)
            
            # Render document
            doc = DocxTemplate(template_stream)
            doc.render(context)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Store in session state
            file_name = f"{sow_num} - {sow_name}.docx"
            st.session_state.file_data = buffer.getvalue()
            st.session_state.generated_file_path = file_name
            
            # Store form data with ALL financial values
            form_data = {
                "option": option,
                "sow_num": sow_num,
                "sow_name": sow_name,
                "Client_Name": Client_Name,
                "start_date": start_date,
                "end_date": end_date,
                "scope_text": scope_text,
                "ser_del": ser_del,
                "pm_client": pm_client,
                "pm_sp": pm_sp,
                "mg_client": mg_client,
                "mg_sp": mg_sp,
                "additional_personnel": additional_personnel,
                "resources_df": resources_df,
                "milestone_df": milestone_df
            }
            
            # Add financial values based on project type
            if option == "T&M":
                form_data["currency_value"] = currency_value
            elif option == "Fixed Fee":
                form_data["Fees_al"] = Fees_al_value
            elif option == "Change Order":
                form_data.update({
                    "Change": st.session_state.get(f"change_{st.session_state.reset_trigger}", "CO-001"),
                    "Fees_co": Fees_co_value,
                    "Fees_sow": Fees_sow_value,
                    "difference": difference_value,
                    "sow_start_date": st.session_state.get(f"sow_start_{st.session_state.reset_trigger}", date.today()),
                    "sow_end_date": st.session_state.get(f"sow_end_{st.session_state.reset_trigger}", date.today())
                })
            
            # Store form_data in session state
            st.session_state.form_data = form_data
            
            # NEW: Create Excel files based on project type
            excel_exporter = ExcelExporter()
            
            if option == "Fixed Fee" and milestone_df is not None and not milestone_df.empty:
                excel_path = excel_exporter.create_fixed_fee_milestone_excel(form_data, milestone_df)
                if excel_path:
                    # Store Excel file data for download
                    with open(excel_path, 'rb') as f:
                        st.session_state.fixed_fee_excel_data = f.read()
                    st.session_state.fixed_fee_excel_name = f"{sow_num}_Milestone_Payments.xlsx"
                    st.info(f"📊 Milestone payment Excel file created: {os.path.basename(excel_path)}")
            
            elif option == "T&M" and resources_df is not None and not resources_df.empty:
                excel_path = excel_exporter.create_tm_resource_excel(form_data, resources_df)
                if excel_path:
                    # Store Excel file data for download
                    with open(excel_path, 'rb') as f:
                        st.session_state.tm_excel_data = f.read()
                    st.session_state.tm_excel_name = f"{sow_num}_Resource_Details.xlsx"
                    st.info(f"📊 Resource details Excel file created: {os.path.basename(excel_path)}")
            
            st.success("✅ SOW document generated successfully!")
            st.session_state.should_increment_on_download = True
            
            # Auto-save to SharePoint
            auto_save_to_sharepoint()
            
            # NEW: Auto-upload if checkbox is checked
            if auto_upload:
                st.info("🔄 Auto-upload to SharePoint enabled - Uploading documents...")
                upload_document_to_sharepoint()
                st.session_state.auto_upload_completed = True
                st.success("✅ Auto-upload to SharePoint completed successfully!")
                st.balloons()
            
        except Exception as e:
            st.error(f"❌ Error generating document: {str(e)}")
            st.exception(e)


def generate_approved_documents(form_data):
    """Generate SOW and Excel documents on approval"""
    try:
        template_manager = TemplateManager()
        excel_exporter = ExcelExporter()
        
        result = {
            "success": True,
            "message": "",
            "document_url": "",
            "excel_url": ""
        }
        
        # Generate SOW document
        context = prepare_document_context(form_data)
        template_stream = template_manager.get_template(form_data.get("option", ""))
        
        doc = DocxTemplate(template_stream)
        doc.render(context)
        
        sow_buffer = BytesIO()
        doc.save(sow_buffer)
        sow_buffer.seek(0)
        
        # Store in session state for upload
        st.session_state.file_data = sow_buffer.getvalue()
        st.session_state.generated_file_path = f"{form_data.get('sow_num', '')} - {form_data.get('sow_name', '')}.docx"
        st.session_state.form_data = form_data
        
        # Generate and store Excel files in session state
        if form_data.get("option") == "Fixed Fee" and form_data.get("milestone_df") is not None:
            milestone_df = form_data.get("milestone_df")
            if not milestone_df.empty:
                excel_path = excel_exporter.create_fixed_fee_milestone_excel(form_data, milestone_df)
                if excel_path and os.path.exists(excel_path):
                    with open(excel_path, 'rb') as f:
                        st.session_state.fixed_fee_excel_data = f.read()
                    st.session_state.fixed_fee_excel_name = f"{form_data.get('sow_num', '')}_Milestone_Payments.xlsx"
                    st.info(f"📊 Milestone payment Excel file created for approval")
        
        elif form_data.get("option") == "T&M" and form_data.get("resources_df") is not None:
            resources_df = form_data.get("resources_df")
            if not resources_df.empty:
                excel_path = excel_exporter.create_tm_resource_excel(form_data, resources_df)
                if excel_path and os.path.exists(excel_path):
                    with open(excel_path, 'rb') as f:
                        st.session_state.tm_excel_data = f.read()
                    st.session_state.tm_excel_name = f"{form_data.get('sow_num', '')}_Resource_Details.xlsx"
                    st.info(f"📊 Resource details Excel file created for approval")
        
        # Upload all documents using the same function
        upload_result = upload_document_to_sharepoint()
        
        if upload_result and upload_result.get("success"):
            result["document_url"] = upload_result.get("word_url", "")
            result["excel_url"] = upload_result.get("excel_url", "")
            result["message"] = "Documents generated and uploaded successfully"
        else:
            result["success"] = False
            result["message"] = upload_result.get("message", "Failed to upload documents")
        
        return result
        
    except Exception as e:
        print(f"❌ Error generating approved documents: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": str(e)
        }

def prepare_document_context(form_data):
    """Prepare document context for template rendering"""
    context = {
        "sow_num": form_data.get("sow_num", ""),
        "sow_name": form_data.get("sow_name", ""),
        "client": form_data.get("Client_Name", ""),
        "project_type": form_data.get("option", ""),
        "start_date": form_data.get("start_date", date.today()).strftime("%B %d, %Y"),
        "end_date": form_data.get("end_date", date.today()).strftime("%B %d, %Y"),
        "generated_date": datetime.now().strftime("%B %d, %Y"),
        "scope_text": form_data.get("scope_text", ""),
        "ser_del": form_data.get("ser_del", ""),
        "pm_client": form_data.get("pm_client", ""),
        "pm_sp": form_data.get("pm_sp", ""),
        "mg_client": form_data.get("mg_client", ""),
        "mg_sp": form_data.get("mg_sp", ""),
        "additional_personnel": form_data.get("additional_personnel", "")
    }
    
    # Add project-specific data
    if form_data.get("option") == "Fixed Fee":
        context["Fees"] = form_data.get("Fees_al", 0)
        if form_data.get("milestone_df") is not None:
            context["milestones"] = form_data["milestone_df"].to_dict(orient="records")
            context["milestone_total"] = form_data["milestone_df"]["net_pay"].sum() if "net_pay" in form_data["milestone_df"].columns else 0
    
    elif form_data.get("option") == "T&M":
        if form_data.get("resources_df") is not None:
            context["resources"] = form_data["resources_df"].to_dict(orient="records")
            context["currency_value"] = form_data["resources_df"]["Estimated $"].sum() if "Estimated $" in form_data["resources_df"].columns else 0
            context["currency_value_str"] = f"${context['currency_value']:,.2f}"
    
    elif form_data.get("option") == "Change Order":
        context.update({
            "Change": form_data.get("Change", ""),
            "Fees_co": form_data.get("Fees_co", 0),
            "Fees_sow": form_data.get("Fees_sow", 0),
            "difference": form_data.get("difference", 0)
        })
    
    return context


def handle_approval_rejection(action):
    """Handle approval/rejection actions - Generate documents on approval"""
    with st.spinner(f"Processing {action.lower()}..."):
        try:
            sharepoint_service = st.session_state.sharepoint_service
            
            # If approving, generate and upload documents first
            if action == Config.STATUS_APPROVED:
                # Collect current form data
                form_data = collect_form_data_from_session()
                
                # Show progress
                with st.spinner("📄 Generating and uploading SOW documents..."):
                    # Generate documents
                    doc_result = generate_approved_documents(form_data)
                    
                    if not doc_result["success"]:
                        st.error(f"Failed to generate documents: {doc_result.get('message', 'Unknown error')}")
                        return False
                    
                    # Update status with document URLs
                    updates = {
                        "status": action,
                        "approver_comments": st.session_state.get("approval_comments", ""),
                        "approved_by": st.session_state.user_email,
                        "approval_date": datetime.now().isoformat(),
                        "DocumentURL": doc_result.get("document_url", ""),
                        "ExcelURL": doc_result.get("excel_url", "")
                    }
                    
                    st.success("✅ Documents generated and uploaded successfully!")
            else:
                # For rejection, just update status
                updates = {
                    "status": action,
                    "approver_comments": st.session_state.get("approval_comments", ""),
                    "rejected_by": st.session_state.user_email,
                    "rejection_date": datetime.now().isoformat()
                }
            
            # Update status in SharePoint
            result = sharepoint_service.update_sow_status(
                item_id=st.session_state.edit_sow_id,
                status=action,
                comments=st.session_state.get("approval_comments", ""),
                approver_email=st.session_state.user_email
            )
            
            if result["success"]:
                if action == Config.STATUS_APPROVED:
                    st.balloons()
                return True
            else:
                st.error(f"Failed to {action.lower()} SOW: {result.get('message', 'Unknown error')}")
                return False
                
        except Exception as e:
            st.error(f"Error: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return False

def auto_save_to_sharepoint():
    """Automatically save SOW data to SharePoint"""
    with st.spinner("💾 Saving SOW data to SharePoint..."):
        try:
            sharepoint_service = st.session_state.sharepoint_service
            form_data = st.session_state.form_data
            
            # Debug form data
            debug_form_data(form_data)
            
            # Prepare SOW record
            sow_record = prepare_sow_data_for_storage(form_data)
            
            # Debug the sow_record
            print(f"🔍 SOW RECORD BEING SENT:")
            print(f"  - TotalValue: {sow_record.get('TotalValue')}")
            print(f"  - Type of TotalValue: {type(sow_record.get('TotalValue'))}")
            print(f"  - Full record keys: {list(sow_record.keys())}")
            
            # Save to SharePoint
            save_result = sharepoint_service.save_sow_record(sow_record)
            
            if save_result["success"]:
                st.session_state.sow_saved = True
                st.session_state.current_sow_data = sow_record
                st.success("✅ SOW data saved to SharePoint!")
                
                # Show the value that was saved
                if sow_record.get('TotalValue', 0) > 0:
                    st.info(f"💵 **Total Value Saved:** ${sow_record['TotalValue']:,.2f}")
                
                # Also save locally as backup
                save_to_local_csv(sow_record)
                
            else:
                st.warning("⚠️ Could not save to SharePoint. Data saved locally only.")
                # Save locally as fallback
                save_to_local_csv(sow_record)
                
        except Exception as e:
            st.error(f"❌ Error saving to SharePoint: {str(e)}")
            # Still try to save locally
            try:
                form_data = st.session_state.form_data
                sow_record = prepare_sow_data_for_storage(form_data)
                save_to_local_csv(sow_record)
                st.info("📋 Data saved locally as backup")
            except:
                pass

def debug_form_data(form_data):
    """Debug function to check form data"""
    print("🔍 DEBUG FORM DATA STRUCTURE:")
    print(f"Keys: {list(form_data.keys())}")
    
    for key, value in form_data.items():
        if key in ['resources_df', 'milestone_df']:
            print(f"  - {key}: {'DataFrame' if value is not None else 'None'}")
            if value is not None:
                print(f"    Shape: {value.shape if hasattr(value, 'shape') else 'N/A'}")
                print(f"    Columns: {list(value.columns) if hasattr(value, 'columns') else 'N/A'}")
        else:
            print(f"  - {key}: {value} (Type: {type(value)})")
    
    # Check for financial values
    print("\n🔍 FINANCIAL VALUES CHECK:")
    print(f"  - Option: {form_data.get('option')}")
    
    if form_data.get('option') == 'T&M':
        print(f"  - currency_value in form_data: {'currency_value' in form_data}")
        if 'currency_value' in form_data:
            print(f"  - currency_value: {form_data['currency_value']}")
    
    elif form_data.get('option') == 'Fixed Fee':
        print(f"  - Fees_al in form_data: {'Fees_al' in form_data}")
        if 'Fees_al' in form_data:
            print(f"  - Fees_al: {form_data['Fees_al']}")
    
    elif form_data.get('option') == 'Change Order':
        print(f"  - difference in form_data: {'difference' in form_data}")
        if 'difference' in form_data:
            print(f"  - difference: {form_data['difference']}")

def show_download_section():
    """Show download and upload options - UPDATED with Excel downloads"""

    # Show auto-upload status
    if st.session_state.get('auto_upload_completed', False):
        st.success("✅ Document has been automatically uploaded to SharePoint!")
        st.balloons()
    
    st.divider()
    st.subheader("📄 Document Ready")
    
    # Create columns based on project type
    if st.session_state.form_data.get('option') == 'Fixed Fee':
        # Fixed Fee: Show Word doc and Excel milestone
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Word Document
            st.markdown("#### 📝 SOW Document")
            st.download_button(
                "📥 Download SOW Document",
                data=st.session_state.file_data,
                file_name=st.session_state.generated_file_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_local_word",
                use_container_width=True
            )
            st.caption("Download the SOW Word document")
        
        with col2:
            # Milestone Excel
            if hasattr(st.session_state, 'fixed_fee_excel_data'):
                st.markdown("#### 📊 Milestone Payments")
                st.download_button(
                    "📥 Download Milestone Excel",
                    data=st.session_state.fixed_fee_excel_data,
                    file_name=st.session_state.fixed_fee_excel_name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_milestone_excel",
                    use_container_width=True
                )
                st.caption("Download milestone payment schedule")
            else:
                st.markdown("#### 📊 Milestone Payments")
                st.info("No milestone data available")
        
        with col3:
            # SharePoint Status
            st.markdown("#### ☁️ SharePoint Status")
            if st.session_state.get('auto_upload_completed', False):
                st.success("✅ Uploaded to SharePoint")
                st.caption("Document stored in 'SOWs' folder")
                if hasattr(st.session_state, 'fixed_fee_excel_data'):
                    st.caption("Excel stored in 'Fixed_Fee_Milestones' folder")
            elif st.session_state.get('document_uploaded', False):
                st.success("✅ Uploaded to SharePoint")
            else:
                st.warning("⚠️ Not uploaded to SharePoint")
                if st.button("📤 Upload to SharePoint Now", use_container_width=True, type="secondary"):
                    upload_document_to_sharepoint()
    
    elif st.session_state.form_data.get('option') == 'T&M':
        # T&M: Show Word doc and Excel resource details
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Word Document
            st.markdown("#### 📝 SOW Document")
            st.download_button(
                "📥 Download SOW Document",
                data=st.session_state.file_data,
                file_name=st.session_state.generated_file_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_local_word",
                use_container_width=True
            )
            st.caption("Download the SOW Word document")
        
        with col2:
            # Resource Excel
            if hasattr(st.session_state, 'tm_excel_data'):
                st.markdown("#### 👥 Resource Details")
                st.download_button(
                    "📥 Download Resource Excel",
                    data=st.session_state.tm_excel_data,
                    file_name=st.session_state.tm_excel_name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_resource_excel",
                    use_container_width=True
                )
                st.caption("Download resource allocation details")
            else:
                st.markdown("#### 👥 Resource Details")
                st.info("No resource data available")
        
        with col3:
            # SharePoint Status
            st.markdown("#### ☁️ SharePoint Status")
            if st.session_state.get('auto_upload_completed', False):
                st.success("✅ Uploaded to SharePoint")
                st.caption("Document stored in 'SOWs' folder")
                if hasattr(st.session_state, 'tm_excel_data'):
                    st.caption("Excel stored in 'TM_Resources' folder")
            elif st.session_state.get('document_uploaded', False):
                st.success("✅ Uploaded to SharePoint")
            else:
                st.warning("⚠️ Not uploaded to SharePoint")
                if st.button("📤 Upload to SharePoint Now", use_container_width=True, type="secondary"):
                    upload_document_to_sharepoint()
    
    else:
        # Other project types (Change Order)
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📝 SOW Document")
            st.download_button(
                "📥 Download SOW Document",
                data=st.session_state.file_data,
                file_name=st.session_state.generated_file_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_local_word",
                use_container_width=True
            )
            st.caption("Download the SOW Word document")
        
        with col2:
            st.markdown("#### ☁️ SharePoint Status")
            if st.session_state.get('auto_upload_completed', False):
                st.success("✅ Uploaded to SharePoint")
                st.caption("Document stored in 'SOWs' folder")
            elif st.session_state.get('document_uploaded', False):
                st.success("✅ Uploaded to SharePoint")
            else:
                st.warning("⚠️ Not uploaded to SharePoint")
                if st.button("📤 Upload to SharePoint Now", use_container_width=True, type="secondary"):
                    upload_document_to_sharepoint()
    
    # Show SOW Details
    if st.session_state.get('sow_saved') and st.session_state.get('current_sow_data'):
        st.divider()
        st.subheader("📋 SOW Details")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("SOW Number", st.session_state.current_sow_data.get("SOWNumber", ""))
            st.metric("Client", st.session_state.current_sow_data.get("Client", ""))
            st.metric("Project Type", st.session_state.current_sow_data.get("ProjectType", ""))
        
        with col2:
            st.metric("Status", st.session_state.current_sow_data.get("Status", ""))
            st.metric("Created By", st.session_state.user_email)
            try:
                total_value = st.session_state.current_sow_data.get("TotalValue", 0)
                if isinstance(total_value, (int, float)):
                    st.metric("Total Value", f"${total_value:,.2f}")
                else:
                    total_value_float = float(str(total_value))
                    st.metric("Total Value", f"${total_value_float:,.2f}")
            except (ValueError, TypeError):
                st.metric("Total Value", str(st.session_state.current_sow_data.get("TotalValue", "0")))

def upload_document_to_sharepoint():
    """Upload documents to SharePoint - UPDATED with separate folders"""
    with st.spinner("Uploading documents to SharePoint..."):
        try:
            sharepoint_service = st.session_state.sharepoint_service
            form_data = st.session_state.form_data
            project_type = form_data.get("option", "")
            
            upload_success = []
            upload_failed = []
            
            # ===== 1. UPLOAD WORD DOCUMENT =====
            st.info(f"📤 Uploading Word document: {st.session_state.generated_file_path}")
            
            word_metadata = {
                "sow_number": form_data.get("sow_num", ""),
                "sow_name": form_data.get("sow_name", ""),
                "client": form_data.get("Client_Name", ""),
                "created_by": st.session_state.user_email,
                "status": Config.STATUS_PENDING,
                "project_type": project_type
            }
            
            # Upload Word document to main SOWs folder
            word_result = sharepoint_service.upload_document(
                st.session_state.file_data,
                st.session_state.generated_file_path,
                word_metadata
            )
            
            if word_result["success"]:
                upload_success.append(f"Word document to 'SOWs' folder")
                st.session_state.word_document_url = word_result.get("data", {}).get("url", "")
            else:
                upload_failed.append(f"Word document: {word_result.get('message', 'Unknown error')}")
            
            # ===== 2. UPLOAD EXCEL FILES =====
            # Determine folder based on project type
            if project_type == "Fixed Fee" and hasattr(st.session_state, 'fixed_fee_excel_data'):
                excel_folder = "Fixed_Fee_Milestones"
                st.info(f"📊 Uploading milestone Excel to '{excel_folder}' folder: {st.session_state.fixed_fee_excel_name}")
                
                excel_metadata = {
                    "sow_number": form_data.get("sow_num", ""),
                    "sow_name": form_data.get("sow_name", ""),
                    "client": form_data.get("Client_Name", ""),
                    "created_by": st.session_state.user_email,
                    "status": Config.STATUS_PENDING,
                    "project_type": project_type,
                    "excel_type": "Milestone Payments"
                }
                
                excel_result = upload_excel_to_sharepoint_folder(
                    sharepoint_service,
                    st.session_state.fixed_fee_excel_data,
                    st.session_state.fixed_fee_excel_name,
                    excel_metadata,
                    excel_folder
                )
                
                if excel_result["success"]:
                    upload_success.append(f"Milestone Excel to '{excel_folder}' folder")
                    st.session_state.excel_document_url = excel_result.get("data", {}).get("url", "")
                else:
                    upload_failed.append(f"Milestone Excel: {excel_result.get('message', 'Unknown error')}")
            
            elif project_type == "T&M" and hasattr(st.session_state, 'tm_excel_data'):
                excel_folder = "TM_Resources"
                st.info(f"👥 Uploading resource Excel to '{excel_folder}' folder: {st.session_state.tm_excel_name}")
                
                excel_metadata = {
                    "sow_number": form_data.get("sow_num", ""),
                    "sow_name": form_data.get("sow_name", ""),
                    "client": form_data.get("Client_Name", ""),
                    "created_by": st.session_state.user_email,
                    "status": Config.STATUS_PENDING,
                    "project_type": project_type,
                    "excel_type": "Resource Details"
                }
                
                excel_result = upload_excel_to_sharepoint_folder(
                    sharepoint_service,
                    st.session_state.tm_excel_data,
                    st.session_state.tm_excel_name,
                    excel_metadata,
                    excel_folder
                )
                
                if excel_result["success"]:
                    upload_success.append(f"Resource Excel to '{excel_folder}' folder")
                    st.session_state.excel_document_url = excel_result.get("data", {}).get("url", "")
                else:
                    upload_failed.append(f"Resource Excel: {excel_result.get('message', 'Unknown error')}")
            
            # ===== 3. SHOW RESULTS =====
            if upload_success:
                st.success("✅ Upload Summary:")
                for success_item in upload_success:
                    st.success(f"   ✓ {success_item}")
                
                st.info("📁 **Files are stored in separate SharePoint folders:**")
                st.info(f"   • Word Document: 'SOWs' folder")
                if project_type == "Fixed Fee" and hasattr(st.session_state, 'fixed_fee_excel_data'):
                    st.info(f"   • Excel File: 'Fixed_Fee_Milestones' folder")
                elif project_type == "T&M" and hasattr(st.session_state, 'tm_excel_data'):
                    st.info(f"   • Excel File: 'TM_Resources' folder")
                
                # Mark auto-upload as completed
                st.session_state.document_uploaded = True
                st.session_state.auto_upload_completed = True
                
                # Return success status with URLs
                return {
                    "success": True,
                    "word_url": st.session_state.get('word_document_url', ''),
                    "excel_url": st.session_state.get('excel_document_url', ''),
                    "message": "Documents uploaded successfully"
                }
            
            if upload_failed:
                st.error("❌ Failed Uploads:")
                for failure in upload_failed:
                    st.error(f"   ✗ {failure}")
                st.session_state.auto_upload_completed = False
                
                return {
                    "success": False,
                    "message": "Some uploads failed"
                }
            
            if upload_success and not upload_failed:
                st.balloons()
                return {
                    "success": True,
                    "message": "All documents uploaded successfully"
                }
                
        except Exception as e:
            st.error(f"❌ Error uploading documents: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            st.session_state.auto_upload_completed = False
            return {
                "success": False,
                "message": str(e)
            }
        

def upload_excel_to_sharepoint_folder(sharepoint_service, file_data, file_name, metadata, folder_name):
    """Upload Excel file to specific SharePoint folder"""
    try:
        # Convert to Base64
        file_base64 = base64.b64encode(file_data).decode('utf-8')
        
        print(f"🔍 DEBUG: Uploading Excel to '{folder_name}' folder")
        print(f"  - File: {file_name}")
        print(f"  - Folder: {folder_name}")
        print(f"  - Metadata: {metadata}")
        
        # Build payload for specific folder
        payload = {
            "operation": "upload_document",
            "library_name": "Onboarding Details",  # Main library
            "folder_path": folder_name,  # Different folder for Excel files
            "file_name": file_name,
            "file_content": file_base64,
            "metadata": {
                "sow_number": metadata.get("sow_number", ""),
                "sow_name": metadata.get("sow_name", ""),
                "client": metadata.get("client", ""),
                "created_by": metadata.get("created_by", ""),
                "status": metadata.get("status", ""),
                "project_type": metadata.get("project_type", ""),
                "excel_type": metadata.get("excel_type", ""),
                "upload_timestamp": datetime.now().isoformat()
            }
        }
        
        # Call Power Automate
        result = sharepoint_service._call_power_automate("upload_document", payload)
        
        if result:
            print(f"✅ SUCCESS: Excel uploaded to '{folder_name}' folder")
            return {
                "success": True,
                "data": result,
                "message": f"Excel uploaded to '{folder_name}' folder successfully"
            }
        
        return {
            "success": False,
            "message": f"No response from Power Automate for '{folder_name}' folder"
        }
                
    except Exception as e:
        print(f"❌ ERROR in upload_excel_to_sharepoint_folder: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "message": f"Fatal error: {str(e)}"
        }

# ============================================================================
# PAGE 2: APPROVAL DASHBOARD
# ============================================================================
def page_approval_dashboard():
    """Approval dashboard for legal team - Simple dataframe with download"""
    
    # Check authorization
    if st.session_state.user_email not in Config.LEGAL_TEAM:
        st.error("⛔ **Access Denied** - This page is for legal team members only.")
        st.info(f"Your email: {st.session_state.user_email}")
        return
    
    st.title("⚖️ Legal Approval Dashboard")
    st.markdown("Review and approve SOW submissions.")
    
    # Check if Get Records flow is configured
    if not Config.POWER_AUTOMATE_URLS["get_records"]:
        st.warning("⚠️ **Get Records Power Automate flow is not configured.**")
        st.info("Please add your Get Records flow URL in the Config class.")
        return
    
    sharepoint_service = st.session_state.sharepoint_service
    
    # ========== FILTERS ==========
    st.subheader("🔍 Filter Options")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        status_filter = st.selectbox(
            "Status",
            ["All", Config.STATUS_PENDING, Config.STATUS_APPROVED, Config.STATUS_REJECTED, Config.STATUS_DRAFT],
            key="approval_status_filter"
        )
    
    with col2:
        client_filter = st.selectbox(
            "Client",
            ["All", "BSC", "Abiomed", "Cognex", "Itaros", "Other"],
            key="approval_client_filter"
        )
    
    with col3:
        project_filter = st.selectbox(
            "Project Type",
            ["All", "Fixed Fee", "T&M", "Change Order"],
            key="approval_project_filter"
        )
    
    # Load records button
    load_clicked = st.button("📥 Load Records", type="primary", use_container_width=True, key="load_records_btn")
    
    if load_clicked or st.session_state.get('sow_dataframe') is not None:
        with st.spinner("Loading SOW records from SharePoint..."):
            # Only fetch new data if button was clicked or no data exists
            if load_clicked or st.session_state.get('sow_dataframe') is None:
                # Build filters
                filters = {}
                if status_filter != "All":
                    filters["status"] = status_filter
                if client_filter != "All":
                    filters["client"] = client_filter
                if project_filter != "All":
                    filters["project_type"] = project_filter
                
                result = sharepoint_service.get_sow_records(**filters)
                
                if not result["success"]:
                    st.error(f"❌ Failed to load records: {result.get('message', 'Unknown error')}")
                    return
                
                df = result["data"]
                
                if df.empty:
                    st.info("📭 No SOW records found with the selected filters.")
                    st.session_state.sow_dataframe = None
                    return
                
                # Store dataframe in session state
                st.session_state.sow_dataframe = df
            else:
                # Use existing dataframe
                df = st.session_state.sow_dataframe
            
            if df.empty:
                st.info("📭 No SOW records found with the selected filters.")
                return
            
            # ========== DISPLAY DATAFRAME WITH VIEW BUTTONS ==========
            st.subheader(f"📋 SOW Records ({len(df)} records)")
            
            # Create a copy for display
            display_df = df.copy()
            
            # Configure column order
            column_order = ['SOWNumber', 'SOWName', 'Client', 'ProjectType', 
                          'Status', 'StartDate', 'EndDate', 'CreatedBy', 'GeneratedDate']
            
            # Filter to only existing columns
            existing_columns = [col for col in column_order if col in display_df.columns]
            
            # Display dataframe
            st.dataframe(
                display_df[existing_columns],
                use_container_width=True,
                height=400,
                column_config={
                    "SOWNumber": st.column_config.TextColumn("SOW #", width="small"),
                    "SOWName": st.column_config.TextColumn("SOW Name", width="medium"),
                    "Client": st.column_config.TextColumn("Client", width="small"),
                    "ProjectType": st.column_config.TextColumn("Type", width="small"),
                    "Status": st.column_config.TextColumn("Status", width="small"),
                    "StartDate": st.column_config.DateColumn("Start Date", format="YYYY-MM-DD", width="small"),
                    "EndDate": st.column_config.DateColumn("End Date", format="YYYY-MM-DD", width="small"),
                    "CreatedBy": st.column_config.TextColumn("Created By", width="medium"),
                    "GeneratedDate": st.column_config.DateColumn("Created Date", format="YYYY-MM-DD", width="small"),
                }
            )
            
            # ========== VIEW SOW FOR APPROVAL ==========
            st.subheader("👁️ View SOW for Approval")
            
            # Create a selectbox with SOW numbers
            sow_options = df['SOWNumber'].unique().tolist()
            
            # Create display names for each SOW
            sow_display_names = []
            for sow in sow_options:
                sow_name = df[df['SOWNumber'] == sow]['SOWName'].iloc[0]
                status = df[df['SOWNumber'] == sow]['Status'].iloc[0]
                display_name = f"{sow} - {sow_name[:30]}{'...' if len(sow_name) > 30 else ''} ({status})"
                sow_display_names.append(display_name)
            
            # Create a dictionary for mapping display names to actual values
            sow_mapping = dict(zip(sow_display_names, sow_options))
            
            # Selectbox for SOW
            selected_display = st.selectbox(
                "Select SOW to review:",
                options=sow_display_names,
                key="sow_select_view",
                index=0
            )
            
            # Get actual SOW number from display name
            selected_sow = sow_mapping[selected_display]
            
            # Find the selected row
            selected_row = df[df['SOWNumber'] == selected_sow].iloc[0]
            
            # View button
            col1, col2 = st.columns([3, 1])
            
            with col1:
                view_btn = st.button(
                    "👁️ View SOW Details & Approve/Reject", 
                    use_container_width=True,
                    type="primary",
                    key="view_sow_btn"
                )
            
            with col2:
                download_btn = st.button(
                    "📥 Download Document", 
                    use_container_width=True,
                    key="download_doc_btn"
                )
            
            if view_btn:
                # Set edit mode and load SOW data
                st.session_state.edit_sow_mode = True
                st.session_state.viewing_for_approval = True
                st.session_state.edit_sow_data = selected_row.to_dict()
                st.session_state.edit_sow_id = selected_row.get('ID')
                st.session_state.edit_mode_enabled = True 
                
                # Load the additional data (resources/milestones) for display
                load_sow_data_for_edit_mode(selected_row.to_dict())
                
                st.rerun()
            
            if download_btn:
                with st.spinner(f"Downloading document for {selected_sow}..."):
                    # Try to get document from SharePoint
                    document_id = selected_row.get('ID')
                    
                    if document_id and Config.POWER_AUTOMATE_URLS["get_document"]:
                        document_bytes = sharepoint_service.get_document(item_id=document_id)
                        
                        if document_bytes:
                            filename = f"{selected_sow}_{selected_row.get('SOWName', 'SOW').replace(' ', '_')}.docx"
                            
                            st.success("✅ Document ready for download!")
                            st.download_button(
                                f"📥 Click to download: {selected_sow}",
                                data=document_bytes,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"download_{selected_sow}"
                            )
                        else:
                            st.warning("⚠️ Document not found in SharePoint.")
                    else:
                        st.error("❌ Document ID not found or get_document flow not configured.")
            
            # ========== EXPORT DATA SECTION ==========
            st.divider()
            st.subheader("📤 Export All Data")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Export to CSV
                csv = df.to_csv(index=False)
                st.download_button(
                    "📥 Download as CSV",
                    data=csv,
                    file_name=f"sow_records_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True,
                    key="export_csv_btn"
                )
            
            with col2:
                # Export to Excel
                excel_buffer = BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='SOW Records')
                
                st.download_button(
                    "📊 Download as Excel",
                    data=excel_buffer.getvalue(),
                    file_name=f"sow_records_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="export_excel_btn"
                )
    
    else:
        # Initial state - show instructions
        st.info("👆 Click 'Load Records' button to load SOW data from SharePoint")
        st.markdown("""
        ### How to use:
        1. Select filters from the dropdowns above
        2. Click **Load Records** button
        3. View all SOW records in the interactive table
        4. Select a SOW from the dropdown below the table
        5. Click **View SOW Details & Approve/Reject** to review and make a decision
        
        ### Features:
        - **View all SOW records** in a clean table
        - **Filter** by status, client, and project type
        - **View SOW details** in read-only mode
        - **Approve or Reject** SOW submissions with comments
        - **Download documents** for review
        """)

# ============================================================================
# PAGE 3: PUBLISHED SOWS
# ============================================================================
# ============================================================================
# PAGE 3: PUBLISHED SOWS (UPDATED WITH BETTER ERROR HANDLING)
# ============================================================================
def page_published_sows():
    """Published SOWs page for all users to view approved SOWs"""
    
    st.title("📚 Published SOWs")
    st.markdown("View all approved SOW documents.")
    
    # Check if Get Records flow is configured
    if not Config.POWER_AUTOMATE_URLS["get_records"]:
        st.warning("⚠️ **Get Records Power Automate flow is not configured.**")
        st.info("Please add your Get Records flow URL in the Config class.")
        return
    
    sharepoint_service = st.session_state.sharepoint_service
    
    # Initialize session state for data persistence
    if 'published_sows_df' not in st.session_state:
        st.session_state.published_sows_df = None
    
    # ========== SIMPLE FILTERS ==========
    st.subheader("🔍 Filter Options")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        status_filter = st.selectbox(
            "Status",
            [Config.STATUS_APPROVED, "All", Config.STATUS_PENDING, Config.STATUS_REJECTED],
            key="published_status_filter"
        )
    
    with col2:
        client_filter = st.selectbox(
            "Client",
            ["All", "BSC", "Abiomed", "Cognex", "Itaros", "Other"],
            key="published_client_filter"
        )
    
    with col3:
        project_filter = st.selectbox(
            "Project Type",
            ["All", "Fixed Fee", "T&M", "Change Order"],
            key="published_project_filter"
        )
    
    # Load records button
    load_clicked = st.button("📥 Load Published SOWs", type="primary", use_container_width=True, key="load_published_btn")
    
    # Data loading logic
    if load_clicked:
        with st.spinner("Loading published SOWs from SharePoint..."):
            # Build filters - Always show approved by default
            filters = {"status": status_filter if status_filter != "All" else Config.STATUS_APPROVED}
            if client_filter != "All":
                filters["client"] = client_filter
            if project_filter != "All":
                filters["project_type"] = project_filter
            
            result = sharepoint_service.get_sow_records(**filters)
            
            if not result["success"]:
                st.error(f"❌ Failed to load records: {result.get('message', 'Unknown error')}")
                st.session_state.published_sows_df = None
                return
            
            df = result["data"]
            
            if df.empty:
                st.info("📭 No published SOWs found with the selected filters.")
                st.session_state.published_sows_df = None
                return
            
            # Store dataframe in session state
            st.session_state.published_sows_df = df
            
    # Display data if loaded
    if st.session_state.published_sows_df is not None:
        df = st.session_state.published_sows_df
        
        if df.empty:
            st.info("📭 No published SOWs found.")
            return
        
        # ========== DISPLAY DATAFRAME ==========
        st.subheader(f"📋 Published SOWs ({len(df)} records)")
        
        # Create a copy for display
        display_df = df.copy()
        
        # Configure column order
        column_order = ['SOWNumber', 'SOWName', 'Client', 'ProjectType', 
                      'Status', 'StartDate', 'EndDate', 'CreatedBy', 'GeneratedDate', 'TotalValue']
        
        # Filter to only existing columns
        existing_columns = [col for col in column_order if col in display_df.columns]
        
        # Display dataframe
        st.dataframe(
            display_df[existing_columns],
            use_container_width=True,
            height=400,
            column_config={
                "SOWNumber": st.column_config.TextColumn("SOW #"),
                "SOWName": st.column_config.TextColumn("SOW Name"),
                "Client": st.column_config.TextColumn("Client"),
                "ProjectType": st.column_config.TextColumn("Type"),
                "Status": st.column_config.TextColumn("Status"),
                "StartDate": st.column_config.DateColumn("Start Date", format="YYYY-MM-DD"),
                "EndDate": st.column_config.DateColumn("End Date", format="YYYY-MM-DD"),
                "CreatedBy": st.column_config.TextColumn("Created By"),
                "GeneratedDate": st.column_config.DateColumn("Created Date", format="YYYY-MM-DD"),
                "TotalValue": st.column_config.NumberColumn("Total Value", format="$%.2f")
            }
        )
        
        # ========== DOWNLOAD SECTION ==========
        st.divider()
        st.subheader("📥 Download Documents")
        
        # Create a selectbox with SOW numbers
        sow_options = df['SOWNumber'].unique().tolist()
        
        # Create display names for each SOW
        sow_display_names = []
        for sow in sow_options:
            sow_name = df[df['SOWNumber'] == sow]['SOWName'].iloc[0]
            project_type = df[df['SOWNumber'] == sow]['ProjectType'].iloc[0]
            display_name = f"{sow} - {sow_name[:30]}{'...' if len(sow_name) > 30 else ''} ({project_type})"
            sow_display_names.append(display_name)
        
        # Create a dictionary for mapping display names to actual values
        sow_mapping = dict(zip(sow_display_names, sow_options))
        
        # Selectbox for SOW
        selected_display = st.selectbox(
            "Select SOW to download:",
            options=sow_display_names,
            key="published_sow_select",
            index=0
        )
        
        # Get actual SOW number from display name
        selected_sow = sow_mapping[selected_display]
        
        # Find the selected row
        selected_row = df[df['SOWNumber'] == selected_sow].iloc[0]
        project_type = selected_row.get('ProjectType', '')
        
        # Debug expander to inspect data structure
        with st.expander("🔍 Debug: View SOW Data Structure", expanded=False):
            st.write("Selected SOW Data:")
            st.json(selected_row.to_dict(), expanded=False)
            
            # Parse and show additional data
            try:
                additional_data = json.loads(selected_row.get("AdditionalData", "{}"))
                st.write("Additional Data Structure:")
                st.json(additional_data, expanded=False)
                
                if project_type == "Fixed Fee":
                    milestones = additional_data.get("project_specific", {}).get("milestones", [])
                    st.write(f"Milestones found: {len(milestones)}")
                    if milestones:
                        st.write("First milestone:", milestones[0] if milestones else "None")
            except:
                st.write("Could not parse AdditionalData")
        
        # Create two columns for download buttons
        col1, col2 = st.columns(2)
        
        with col1:
            # Download SOW Document button
            download_sow_clicked = st.button(
                "📄 Download SOW Document", 
                use_container_width=True,
                type="primary",
                key="download_sow_btn"
            )
        
        with col2:
            # Download Calculation Sheet button (only show for Fixed Fee and T&M)
            if project_type in ["Fixed Fee", "T&M"]:
                calculation_sheet_label = "📊 Download Milestone Sheet" if project_type == "Fixed Fee" else "📊 Download Resource Sheet"
                download_calc_clicked = st.button(
                    calculation_sheet_label, 
                    use_container_width=True,
                    type="secondary",
                    key="download_calc_btn"
                )
            else:
                st.info("ℹ️ No calculation sheet available for Change Order")
                download_calc_clicked = False
        
        # Handle SOW document download
        if download_sow_clicked:
            with st.spinner(f"Fetching SOW document for {selected_sow}..."):
                # Try to get document from SharePoint using the flow
                document_id = selected_row.get('ID')
                
                if document_id and Config.POWER_AUTOMATE_URLS["get_document"]:
                    document_bytes = sharepoint_service.get_document(item_id=document_id)
                    
                    if document_bytes:
                        # Create download button immediately
                        filename = f"{selected_sow}_{selected_row.get('SOWName', 'SOW').replace(' ', '_')}.docx"
                        
                        st.success("✅ Document ready for download!")
                        st.download_button(
                            f"📥 Click to download: {selected_sow}",
                            data=document_bytes,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=f"download_sow_{selected_sow}"
                        )
                    else:
                        st.warning("⚠️ Document not found in SharePoint.")
                        # Create a fallback document
                        from docx import Document
                        doc = Document()
                        doc.add_heading(f"SOW: {selected_sow}", 0)
                        doc.add_paragraph(f"SOW Name: {selected_row.get('SOWName', 'N/A')}")
                        doc.add_paragraph(f"Client: {selected_row.get('Client', 'N/A')}")
                        doc.add_paragraph(f"Project Type: {selected_row.get('ProjectType', 'N/A')}")
                        doc.add_paragraph(f"Status: {selected_row.get('Status', 'N/A')}")
                        doc.add_paragraph(f"Created By: {selected_row.get('CreatedBy', 'N/A')}")
                        doc.add_paragraph(f"Created Date: {selected_row.get('GeneratedDate', 'N/A')}")
                        
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.download_button(
                            "📝 Download Summary Document",
                            data=buffer.getvalue(),
                            file_name=f"{selected_sow}_Summary.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                else:
                    st.error("❌ Document ID not found or get_document flow not configured.")
        
        # Handle calculation sheet download
        if download_calc_clicked:
            with st.spinner(f"Generating calculation sheet for {selected_sow}..."):
                try:
                    # Parse additional data to get milestone/resource information
                    additional_data = {}
                    try:
                        additional_data = json.loads(selected_row.get("AdditionalData", "{}"))
                        st.write("Debug - Additional Data loaded successfully")
                    except Exception as e:
                        st.error(f"Error parsing AdditionalData: {str(e)}")
                        additional_data = {}
                    
                    project_specific = additional_data.get("project_specific", {})
                    
                    # Debug output
                    st.write(f"Debug - Project Type: {project_type}")
                    st.write(f"Debug - Project Specific Keys: {list(project_specific.keys())}")
                    
                    # Create form data structure needed for Excel generation
                    form_data = {
                        "sow_num": selected_sow,
                        "sow_name": selected_row.get("SOWName", ""),
                        "Client_Name": selected_row.get("Client", ""),
                        "option": project_type,
                        "start_date": selected_row.get("StartDate", date.today()),
                        "end_date": selected_row.get("EndDate", date.today())
                    }
                    
                    excel_exporter = ExcelExporter()
                    
                    if project_type == "Fixed Fee":
                        # Get milestone data
                        milestones_list = project_specific.get("milestones", [])
                        st.write(f"Debug - Milestones list length: {len(milestones_list)}")
                        
                        if milestones_list and len(milestones_list) > 0:
                            # Convert to DataFrame
                            milestone_df = pd.DataFrame(milestones_list)
                            st.write(f"Debug - Milestone DataFrame shape: {milestone_df.shape}")
                            st.write(f"Debug - Milestone DataFrame columns: {list(milestone_df.columns)}")
                            
                            # Check if required columns exist
                            required_cols = ['milestone_no', 'services', 'due_date', 'allocation', 'net_pay']
                            missing_cols = [col for col in required_cols if col not in milestone_df.columns]
                            
                            if missing_cols:
                                st.warning(f"⚠️ Missing columns in milestone data: {missing_cols}")
                                st.write("Available columns:", list(milestone_df.columns))
                                
                                # Try to map columns if they have different names
                                column_mapping = {
                                    'milestone_no': ['milestone_no', 'Milestone #', 'Milestone No', 'Milestone Number'],
                                    'services': ['services', 'Services', 'Services / Deliverables', 'Deliverables'],
                                    'due_date': ['due_date', 'Due Date', 'Milestone Due Date', 'Date'],
                                    'allocation': ['allocation', 'Allocation', 'Payment Allocation (%)', 'Allocation %'],
                                    'net_pay': ['net_pay', 'Net Pay', 'Payment Amount ($)', 'Payment']
                                }
                                
                                for req_col, possible_names in column_mapping.items():
                                    if req_col not in milestone_df.columns:
                                        for possible_name in possible_names:
                                            if possible_name in milestone_df.columns:
                                                milestone_df[req_col] = milestone_df[possible_name]
                                                st.write(f"Mapped {possible_name} to {req_col}")
                                                break
                            
                            # Add Fees_al if available
                            form_data["Fees_al"] = project_specific.get("fees", 0)
                            st.write(f"Debug - Fees_al: {form_data['Fees_al']}")
                            
                            # Generate Excel
                            excel_path = excel_exporter.create_fixed_fee_milestone_excel(form_data, milestone_df)
                            st.write(f"Debug - Excel path: {excel_path}")
                            
                            if excel_path and os.path.exists(excel_path):
                                with open(excel_path, 'rb') as f:
                                    excel_data = f.read()
                                
                                excel_filename = f"{selected_sow}_Milestone_Payments.xlsx"
                                
                                st.success("✅ Milestone sheet generated successfully!")
                                st.download_button(
                                    "📥 Download Milestone Sheet",
                                    data=excel_data,
                                    file_name=excel_filename,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True,
                                    key=f"download_milestone_{selected_sow}"
                                )
                                
                                # Clean up temp file
                                try:
                                    os.remove(excel_path)
                                except:
                                    pass
                            else:
                                st.error("❌ Failed to generate milestone sheet - Excel creation failed")
                                st.write("Debug - Check if ExcelExporter.create_fixed_fee_milestone_excel returned None")
                        else:
                            st.warning("⚠️ No milestone data found for this SOW")
                            st.write("Debug - Milestones list is empty or None")
                    
                    elif project_type == "T&M":
                        # Get resource data
                        resources_list = project_specific.get("resources", [])
                        st.write(f"Debug - Resources list length: {len(resources_list)}")
                        
                        if resources_list and len(resources_list) > 0:
                            # Convert to DataFrame
                            resources_df = pd.DataFrame(resources_list)
                            st.write(f"Debug - Resource DataFrame shape: {resources_df.shape}")
                            st.write(f"Debug - Resource DataFrame columns: {list(resources_df.columns)}")
                            
                            # Check if required columns exist
                            required_cols = ['Role', 'Location', 'Start Date', 'End Date', 'Allocation %', 'Hrs/Day', 'Rate/hr ($)', 'Estimated $']
                            missing_cols = [col for col in required_cols if col not in resources_df.columns]
                            
                            if missing_cols:
                                st.warning(f"⚠️ Missing columns in resource data: {missing_cols}")
                                st.write("Available columns:", list(resources_df.columns))
                            
                            # Generate Excel
                            excel_path = excel_exporter.create_tm_resource_excel(form_data, resources_df)
                            st.write(f"Debug - Excel path: {excel_path}")
                            
                            if excel_path and os.path.exists(excel_path):
                                with open(excel_path, 'rb') as f:
                                    excel_data = f.read()
                                
                                excel_filename = f"{selected_sow}_Resource_Details.xlsx"
                                
                                st.success("✅ Resource sheet generated successfully!")
                                st.download_button(
                                    "📥 Download Resource Sheet",
                                    data=excel_data,
                                    file_name=excel_filename,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True,
                                    key=f"download_resource_{selected_sow}"
                                )
                                
                                # Clean up temp file
                                try:
                                    os.remove(excel_path)
                                except:
                                    pass
                            else:
                                st.error("❌ Failed to generate resource sheet - Excel creation failed")
                        else:
                            st.warning("⚠️ No resource data found for this SOW")
                    
                except Exception as e:
                    st.error(f"❌ Error generating calculation sheet: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())
        
        # ========== EXPORT DATA SECTION ==========
        st.divider()
        st.subheader("📤 Export All Published Data")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Export to CSV
            csv = df.to_csv(index=False)
            st.download_button(
                "📥 Download as CSV",
                data=csv,
                file_name=f"published_sows_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True,
                key="export_published_csv_btn"
            )
        
        with col2:
            # Export to Excel
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Published SOWs')
            
            st.download_button(
                "📊 Download as Excel",
                data=excel_buffer.getvalue(),
                file_name=f"published_sows_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="export_published_excel_btn"
            )
    
    elif not st.session_state.get('published_sows_df') and not load_clicked:
        # Initial state - show instructions
        st.info("👆 Click 'Load Published SOWs' button to view all approved SOW documents")
        st.markdown("""
        ### How to use:
        1. Select filters from the dropdowns above
        2. Click **Load Published SOWs** button
        3. View all published SOW records in the interactive dataframe
        4. Select a SOW number from the dropdown below the table
        5. Click **Download SOW Document** to download the main SOW document
        6. Click **Download Milestone Sheet** (for Fixed Fee) or **Download Resource Sheet** (for T&M) to download calculation sheets
        
        ### Features:
        - **View all published SOW records** (approved status)
        - **Filter** by client and project type
        - **Download individual SOW documents** by selecting from the dropdown
        - **Download calculation sheets** (milestone or resource details) for approved SOWs
        - **Export all published data** as CSV or Excel
        
        ### Note:
        - Only approved SOWs are shown by default
        - Calculation sheets are available for Fixed Fee and T&M projects
        - All users can access this page to view published SOWs
        """)
    
    # Add a reset button at the bottom
    st.divider()
    if st.button("🔄 Clear & Refresh", type="secondary", use_container_width=True):
        st.session_state.published_sows_df = None
        st.rerun()

# ============================================================================
# MAIN APPLICATION
# ============================================================================
def main():
    """Main application entry point"""
    
    # Initialize session state
    init_session_state()
    
    # Check authentication
    if not st.session_state.is_authenticated:
        login_page()
        return
    
    # Render header with user info and logout button
    render_header()
    
    # Navigation
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Show different pages based on user role and edit mode
    if st.session_state.edit_sow_mode:
        # If in edit mode, always show SOW Generator page
        page_sow_generator()
    else:
        # Normal navigation
        if st.session_state.user_role == 'legal':
            pages = ["SOW Generator", "Approval Dashboard", "Published SOWs"]
        else:
            pages = ["SOW Generator", "Published SOWs"]
        
        # Navigation tabs
        tab = st.radio(
            "Select Page:",
            pages,
            horizontal=True,
            label_visibility="collapsed"
        )
        
        # Page routing
        if tab == "SOW Generator":
            page_sow_generator()
        elif tab == "Approval Dashboard":
            page_approval_dashboard()
        elif tab == "Published SOWs":
            page_published_sows()

# ============================================================================
# RUN APPLICATION
# ============================================================================
if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"🚨 Application error: {str(e)}")
        st.info("Please check the Power Automate configuration in the Config class.")
        st.button("🔄 Restart Application", on_click=st.rerun())

